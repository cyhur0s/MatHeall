from docx import Document
import sys
sys.stdout.reconfigure(encoding="utf-8")
p = r"D:\KULIAH\Penelitian Ilmiah\blackbox.docx"
d = Document(p)
print("P", len(d.paragraphs), "T", len(d.tables))
for i, x in enumerate(d.paragraphs, 1):
    if x.text.strip(): print(f"P{i:03d}|{x.style.name if x.style else '[none]'}|{x.text.strip()}")
for i, table in enumerate(d.tables, 1):
    print(f"TABLE {i}: {len(table.rows)}x{len(table.columns)}")
    for row in table.rows:
        print(" | ".join(cell.text.replace("\n", " / ") for cell in row.cells))
