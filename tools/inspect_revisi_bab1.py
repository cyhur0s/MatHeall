from docx import Document
from zipfile import ZipFile
from pathlib import Path
import sys

sys.stdout.reconfigure(encoding="utf-8")
path = Path(r"C:\Users\HP\Downloads\Revisi_Bab1_dan_MetodePenelitian_MatHeal.docx")
doc = Document(path)
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
for i, p in enumerate(doc.paragraphs, 1):
    if p.text.strip():
        style = p.style.name if p.style else "[no style]"
        print(f"P{i:03d}|{style}|{p.text.strip()}")
for i, table in enumerate(doc.tables, 1):
    print(f"TABLE {i} {len(table.rows)}x{len(table.columns)}")
    for row in table.rows:
        print(" | ".join(cell.text.replace("\n", " / ") for cell in row.cells))
with ZipFile(path) as z:
    names = z.namelist()
    print("comments=", "word/comments.xml" in names, "track_changes=", "word/people.xml" in names or any("tracked" in n for n in names))
    if "word/comments.xml" in names:
        print(z.read("word/comments.xml").decode("utf-8", errors="replace"))
