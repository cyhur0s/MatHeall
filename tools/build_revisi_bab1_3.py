from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parents[1] / "docs" / "Revisi_Terintegrasi_Bab_I_III_MatHeal.docx"


def set_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size); run.bold = bold; run.italic = italic


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill); tcpr.append(shd)


def page_no(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(); set_font(r, 10)
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = "PAGE"
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    r._r.extend([b, i, e])


def setup(doc):
    sec = doc.sections[0]
    sec.top_margin, sec.bottom_margin, sec.left_margin, sec.right_margin = Cm(4), Cm(3), Cm(4), Cm(3)
    sec.header_distance, sec.footer_distance = Cm(1.5), Cm(1.5)
    for name, size in (("Normal",12),("Heading 1",14),("Heading 2",12),("Heading 3",12)):
        s=doc.styles[name]; s.font.name="Times New Roman"; s._element.rPr.rFonts.set(qn("w:ascii"),"Times New Roman"); s._element.rPr.rFonts.set(qn("w:hAnsi"),"Times New Roman"); s.font.size=Pt(size)
        if name != "Normal": s.font.bold=True
    doc.styles["Normal"].paragraph_format.line_spacing=1.5
    page_no(sec.footer.paragraphs[0])


def title(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(18)
    set_font(p.add_run(text),14,True)


def h(doc, text, level=2):
    p=doc.add_paragraph(style=f"Heading {level}"); p.paragraph_format.keep_with_next=True; p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(6)
    set_font(p.add_run(text), 12 if level > 1 else 14, True); return p


def para(doc, text, first=True):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing=1.5
    if first: p.paragraph_format.first_line_indent=Cm(1.25)
    set_font(p.add_run(text)); return p


def bullet(doc, text, n=None):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing=1.5; p.paragraph_format.left_indent=Cm(.8); p.paragraph_format.first_line_indent=Cm(-.8)
    prefix=f"{n}. " if n is not None else "• "
    set_font(p.add_run(prefix),12,True); set_font(p.add_run(text)); return p


def table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j, head in enumerate(headers):
        c=t.rows[0].cells[j]; shade(c,"397582"); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(p.add_run(head),10,True)
        for run in p.runs: run.font.color.rgb=RGBColor(255,255,255)
    for row in rows:
        cells=t.add_row().cells
        for j, val in enumerate(row):
            p=cells[j].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT; set_font(p.add_run(str(val)),9)
            cells[j].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths: cells[j].width=Inches(widths[j])
    doc.add_paragraph()
    return t


def chapter1(doc):
    title(doc,"BAB I\nPENDAHULUAN")
    h(doc,"1.1 Latar Belakang")
    for text in [
        "Matematika informatika merupakan fondasi bagi kompetensi komputasional, seperti pemahaman logika, himpunan, relasi, fungsi, bilangan biner, matriks, peluang, dan graf. Penguasaan konsep tersebut membantu mahasiswa memahami algoritma, struktur data, basis data, dan mata kuliah komputasi lanjutan. Oleh karena itu, proses belajar memerlukan materi terstruktur, contoh bertahap, latihan, serta umpan balik terhadap kesalahan pengguna.",
        "Media e-learning yang hanya menyediakan dokumen atau kuis dengan pencocokan jawaban akhir belum sepenuhnya mendukung kebutuhan tersebut. Dalam matematika, jawaban yang benar dapat ditulis dengan notasi atau metode yang berbeda. Pengguna memerlukan penjelasan konsep dan langkah penyelesaian, bukan hanya informasi benar atau salah. Kondisi ini menunjukkan perlunya media yang menghubungkan materi, latihan bertingkat, umpan balik, dan pemantauan progres dalam satu alur belajar.",
        "Gamifikasi dapat digunakan untuk mendukung konsistensi latihan melalui level, poin, streak, heart, misi harian, indikator progres, dan sertifikat. Elemen ini perlu ditempatkan sebagai dukungan terhadap aktivitas belajar dan bukan sekadar dekorasi antarmuka (Deterding dkk., 2011; Sailer dan Homner, 2020). Di sisi lain, AI generatif dapat dimanfaatkan sebagai pendamping tanya jawab dan pemberi umpan balik esai, dengan batasan bahwa AI bukan sumber kebenaran akademik tunggal.",
        "Berdasarkan kebutuhan tersebut, dikembangkan MatHeal, yaitu aplikasi e-learning matematika informatika berbasis web yang menyediakan modul materi, kuis bertingkat, gamifikasi, progres, sertifikat, dashboard admin, AI Tutor AskMatheal, dan penilaian esai berbasis rubrik. Penelitian ini tidak melibatkan responden mahasiswa, kuesioner SUS, atau validator dosen. Oleh karena itu, hasil penelitian dibatasi pada pengembangan aplikasi dan pengujian fungsional black-box yang dilakukan oleh peneliti; penelitian ini tidak menyimpulkan usability pengguna, akurasi akademik AI, maupun peningkatan hasil belajar."
    ]: para(doc,text)
    h(doc,"1.2 Rumusan Masalah")
    para(doc,"Rumusan masalah penelitian ini adalah sebagai berikut.",False)
    for i,t in enumerate([
        "Bagaimana merancang dan mengembangkan aplikasi e-learning matematika informatika berbasis web yang mengintegrasikan materi, kuis bertingkat, progres, sertifikat, dan dashboard administrator?",
        "Bagaimana mengintegrasikan AI Tutor Google Gemini sebagai media tanya jawab matematika dan penilai jawaban esai berbasis rubrik?",
        "Bagaimana menerapkan gamifikasi berupa level, poin, streak, heart, misi harian, progres, dan sertifikat untuk mendukung keteraturan belajar pengguna?",
        "Bagaimana hasil pengujian fungsional black-box terhadap fitur utama MatHeal yang dilakukan oleh peneliti?"
    ],1): bullet(doc,t,i)
    h(doc,"1.3 Batasan Masalah")
    for i,t in enumerate([
        "Materi dibatasi pada 17 topik matematika informatika yang tersedia pada aplikasi.",
        "Aplikasi memiliki peran user dan admin; user menggunakan layanan belajar, sedangkan admin mengelola materi, soal, pengguna, dan aktivitas.",
        "Kuis terdiri atas level mudah, sedang, dan sulit. Level berikutnya terbuka setelah level sebelumnya lulus dengan nilai minimal 70 persen.",
        "AI Tutor menggunakan Google Gemini sebagai pendamping belajar dan umpan balik otomatis, bukan pengganti pengajar atau keputusan akademik resmi.",
        "Sertifikat hanya menjadi bukti penyelesaian materi di dalam aplikasi.",
        "Penelitian tidak menggunakan SUS, responden, validasi dosen, pengujian akurasi AI terhadap ahli, maupun eksperimen pretest-posttest.",
        "Pengujian yang dilakukan adalah black-box testing fungsional oleh peneliti pada fitur utama MatHeal."
    ],1): bullet(doc,t,i)
    h(doc,"1.4 Tujuan Penelitian")
    for i,t in enumerate([
        "Mengembangkan aplikasi e-learning matematika informatika berbasis web bernama MatHeal yang menghubungkan materi, kuis bertingkat, progres, sertifikat, dan dashboard administrator.",
        "Mengintegrasikan AI Tutor Google Gemini untuk memberi penjelasan matematika dan menilai jawaban esai berdasarkan ketepatan jawaban akhir serta validitas metode.",
        "Menerapkan gamifikasi berupa level, poin, streak, heart, misi harian, progres, dan sertifikat.",
        "Menguji kesesuaian fungsi utama MatHeal melalui black-box testing yang dilakukan oleh peneliti."
    ],1): bullet(doc,t,i)
    h(doc,"1.5 Manfaat Penelitian")
    for i,t in enumerate([
        "Bagi pengguna, MatHeal menyediakan materi, latihan bertingkat, umpan balik, dan informasi progres dalam satu aplikasi.",
        "Bagi administrator, aplikasi menyediakan fasilitas pengelolaan materi, bank soal, pengguna, dan aktivitas aplikasi.",
        "Bagi penelitian selanjutnya, dokumentasi aplikasi dan hasil black-box dapat menjadi dasar untuk evaluasi usability, validasi ahli, dan pengukuran efektivitas pembelajaran."
    ],1): bullet(doc,t,i)
    h(doc,"1.6 Metode Penelitian")
    para(doc,"Penelitian ini menggunakan pendekatan Research and Development dengan model prototyping. Tahapannya meliputi identifikasi kebutuhan, perancangan cepat, pembangunan prototipe, evaluasi internal, penyempurnaan, dan pengujian fungsional. Data pengembangan diperoleh melalui studi pustaka, analisis kebutuhan aplikasi, dokumentasi kode, dan hasil black-box testing. Pengujian dilakukan oleh peneliti dengan membandingkan perilaku aplikasi terhadap hasil yang diharapkan tanpa meninjau struktur internal kode (Pressman dan Maxim, 2020).")
    h(doc,"1.7 Sistematika Penulisan")
    for k,t in [("BAB I","Pendahuluan, berisi latar belakang, rumusan masalah, batasan, tujuan, manfaat, metode, dan sistematika penulisan."),("BAB II","Landasan Teori, berisi teori e-learning, matematika informatika, gamifikasi, AI Tutor, penilaian esai, teknologi web, dan black-box testing."),("BAB III","Analisis, Perancangan, Implementasi, dan Pengujian, berisi kebutuhan sistem, rancangan, implementasi, serta hasil black-box testing."),("BAB IV","Penutup, berisi kesimpulan, keterbatasan, dan saran.")]:
        bullet(doc,t.replace(t,t),None); doc.paragraphs[-1].clear(); set_font(doc.paragraphs[-1].add_run(k+" "),12,True); set_font(doc.paragraphs[-1].add_run(t))


def chapter2(doc):
    doc.add_page_break(); title(doc,"BAB II\nLANDASAN TEORI")
    h(doc,"2.1 E-Learning")
    para(doc,"E-learning merupakan pembelajaran yang memanfaatkan teknologi digital untuk menyediakan materi, interaksi, latihan, dan evaluasi. Dalam MatHeal, e-learning diposisikan sebagai lingkungan belajar yang menggabungkan modul, kuis, umpan balik, dan pencatatan progres, bukan hanya repositori berkas.")
    h(doc,"2.2 Matematika Informatika")
    para(doc,"Matematika informatika mencakup konsep matematika yang relevan bagi ilmu komputer, antara lain logika, himpunan, relasi, fungsi, bilangan biner, matriks, peluang, dan graf. Konsep tersebut membutuhkan latihan bertahap karena kesalahan pemahaman konsep dasar dapat memengaruhi pemahaman topik berikutnya.")
    h(doc,"2.3 Gamifikasi dalam Pembelajaran")
    para(doc,"Gamifikasi adalah penerapan elemen desain permainan pada konteks nonpermainan (Deterding dkk., 2011). Dalam pembelajaran, level, poin, progres, streak, misi, dan penghargaan dapat memberi tujuan jangka pendek dan umpan balik atas aktivitas belajar. Efektivitasnya dipengaruhi konteks dan rancangan, sehingga elemen gamifikasi MatHeal dihubungkan dengan penyelesaian kuis dan progres belajar, bukan kompetisi antarpengguna.")
    h(doc,"2.4 Artificial Intelligence, AI Tutor, dan Google Gemini")
    para(doc,"AI Tutor adalah sistem yang menyediakan bantuan pembelajaran melalui interaksi bahasa alami. AskMatheal menggunakan Google Gemini melalui layanan backend untuk menjawab pertanyaan matematika dan memberikan umpan balik. Pemanfaatannya dibatasi sebagai pendamping belajar; pengguna tetap perlu memeriksa penjelasan dan hasil AI dengan materi serta aturan pembelajaran yang berlaku.")
    h(doc,"2.5 Penilaian Jawaban Esai")
    para(doc,"Penilaian esai matematika tidak hanya memperhatikan jawaban akhir, tetapi juga validitas konsep, metode, dan konsistensi langkah. MatHeal memakai pemeriksaan deterministik untuk kecocokan jawaban atau nilai numerik ekuivalen, kemudian menggunakan Gemini untuk penilaian semantik bila diperlukan. Rubrik terdiri atas jawaban akhir 50 poin, validitas metode 40 poin, dan konsistensi proses 10 poin. Hasil fitur ini hanya diuji secara fungsional, bukan divalidasi terhadap penilaian ahli pada penelitian ini.")
    h(doc,"2.6 Aplikasi Web dan Arsitektur Client-Server")
    para(doc,"MatHeal menggunakan frontend React dan Vite, backend PHP/MySQL, serta gateway AI Node.js/Express. Frontend mengelola antarmuka pengguna. Backend PHP mengelola autentikasi, materi, soal, aktivitas, dan progres. Gateway Node.js meneruskan permintaan AskMatheal serta penilaian esai ke Google Gemini dan menjaga kunci API tetap berada di server.")
    h(doc,"2.7 Black-Box Testing")
    para(doc,"Black-box testing adalah pengujian perangkat lunak dengan memeriksa kesesuaian masukan dan keluaran tanpa menganalisis struktur internal kode. Peneliti menentukan skenario, menjalankan fitur melalui antarmuka atau endpoint, mencatat hasil aktual, lalu menetapkan status berhasil apabila hasil aktual sesuai dengan hasil yang diharapkan. Pengujian ini sesuai untuk memeriksa fungsi utama MatHeal secara mandiri oleh peneliti.")
    h(doc,"2.8 Kerangka Pemikiran")
    para(doc,"Masalah berupa materi dan latihan yang belum terhubung, kebutuhan umpan balik, serta perlunya konsistensi latihan dijawab melalui MatHeal. Sistem mengintegrasikan modul materi, kuis bertingkat, gamifikasi, progres, AI Tutor, dan penilaian esai. Keluaran penelitian adalah aplikasi MatHeal beserta bukti kesesuaian fungsi utama melalui black-box testing. Penelitian tidak menarik kesimpulan mengenai peningkatan hasil belajar, usability pengguna, atau akurasi akademik AI karena ketiga aspek tersebut membutuhkan rancangan evaluasi yang berbeda.")


def chapter3(doc):
    doc.add_page_break(); title(doc,"BAB III\nANALISIS, PERANCANGAN, IMPLEMENTASI, DAN PENGUJIAN")
    h(doc,"3.1 Jenis dan Pendekatan Model Penelitian")
    para(doc,"Penelitian ini merupakan penelitian pengembangan perangkat lunak dengan model prototyping. Model ini digunakan karena fitur dan antarmuka dapat dibangun secara bertahap, diperiksa secara internal, lalu disempurnakan. Pengujian fungsional dilakukan oleh peneliti menggunakan black-box testing.")
    h(doc,"3.2 Objek, Subjek, dan Lokasi Penelitian")
    para(doc,"Objek penelitian adalah aplikasi MatHeal. Penelitian tidak melibatkan subjek manusia. Subjek pengujian berupa fungsi aplikasi, yaitu autentikasi, materi, kuis, progres, gamifikasi, sertifikat, AI Tutor, dan dashboard admin. Lokasi pengembangan adalah Universitas Gunadarma, Bekasi Utara, Jawa Barat. Periode penelitian disesuaikan dengan waktu pengembangan dan pengujian yang sebenarnya.")
    h(doc,"3.3 Teknik Pengumpulan Data")
    for i,t in enumerate(["Studi pustaka mengenai e-learning, gamifikasi, AI Tutor, penilaian esai, dan pengujian perangkat lunak.","Analisis kebutuhan dan dokumentasi sistem, meliputi kode, struktur basis data, konfigurasi, modul materi, serta bank soal.","Black-box testing oleh peneliti terhadap fungsi utama aplikasi."] ,1): bullet(doc,t,i)
    h(doc,"3.4 Tahapan Pengembangan Prototipe")
    table(doc,["Tahap","Aktivitas","Keluaran"],[
        ("Requirements gathering","Mengidentifikasi masalah, batasan, dan kebutuhan user/admin melalui studi pustaka serta analisis aplikasi.","Daftar kebutuhan sistem."),
        ("Perancangan cepat","Menyusun arsitektur, alur belajar, aturan gamifikasi, rancangan data, dan antarmuka.","Rancangan sistem dan model data."),
        ("Pembangunan prototipe","Mengimplementasikan frontend, backend PHP/MySQL, dan layanan AI Node.js/Express.","Prototipe fungsional."),
        ("Evaluasi internal","Memeriksa navigasi, kuis, progres, gamifikasi, dan penanganan gangguan AI.","Daftar temuan."),
        ("Penyempurnaan","Memperbaiki fungsi, antarmuka, sinkronisasi progres, dan keamanan.","Versi kandidat uji."),
        ("Pengujian fungsional","Menjalankan skenario black-box pada fungsi utama.","Hasil pengujian fungsional.")
    ],[1.25,3.6,2.1])
    h(doc,"3.5 Analisis Permasalahan")
    para(doc,"Permasalahan yang menjadi dasar pengembangan adalah materi dan latihan yang belum terhubung dalam satu alur, umpan balik yang terbatas pada jawaban benar atau salah, kesulitan memantau progres, dan kebutuhan latihan yang konsisten. MatHeal menjawab masalah tersebut dengan modul materi, kuis bertingkat, umpan balik, progres, gamifikasi, dan AI Tutor sebagai pendamping belajar.")
    h(doc,"3.6 Analisis Kebutuhan Pengguna")
    h(doc,"3.6.1 Kebutuhan Fungsional User",3)
    for t in ["Registrasi, login, logout, dan permintaan reset password.","Akses modul materi dan kuis bertingkat.","Pengerjaan soal pilihan ganda, benar-salah, dan esai; penerimaan feedback serta hasil kuis.","Pemantauan progres, poin, streak, heart, misi harian, sertifikat, dan catatan belajar.","Akses AskMatheal dan penyimpanan riwayat percakapan."]: bullet(doc,t)
    h(doc,"3.6.2 Kebutuhan Fungsional Admin",3)
    for t in ["Melihat ringkasan pengguna, aktivitas, materi dibuka, dan kuis diselesaikan.","Mengelola pengguna, materi PDF, bank soal, video panduan, dan data pengembang.","Melihat aktivitas operasional yang diperlukan dashboard."]: bullet(doc,t)
    h(doc,"3.6.3 Kebutuhan Nonfungsional",3)
    table(doc,["Aspek","Kebutuhan"],[
        ("Usability","Navigasi dan feedback menggunakan istilah yang konsisten."),("Kompatibilitas","Tampilan menyesuaikan mobile, tablet, laptop, dan desktop."),("Keamanan","Password di-hash, token dan role diverifikasi, kunci AI berada di server."),("Keandalan","Gangguan AI ditangani tanpa merusak sesi kuis."),("Maintainability","Frontend, API, gateway AI, katalog kuis, dan utilitas dipisahkan.")
    ],[1.6,5.4])
    h(doc,"3.7 Arsitektur Sistem")
    para(doc,"Arsitektur MatHeal terdiri atas frontend React, backend PHP/MySQL, dan gateway AI Node.js/Express. Frontend mengirim permintaan data akun, materi, soal, aktivitas, dan progres kepada PHP API. Data disimpan dalam MySQL. Pertanyaan AskMatheal dan penilaian esai dikirim ke gateway AI, kemudian gateway menghubungi Google Gemini. Token pengguna diverifikasi melalui backend PHP sebelum endpoint AI diproses.")
    h(doc,"3.8 Perancangan Alur Belajar")
    para(doc,"Pengguna memilih materi, mempelajari modul, lalu mengerjakan kuis pada level yang terbuka. Level sedang terbuka setelah level mudah lulus dan level sulit terbuka setelah level sedang lulus. Satu kuis berisi 10 soal. Nilai minimal 7 jawaban benar dari 10 soal menghasilkan status lulus, progres tersimpan, dan level berikutnya dapat dibuka. Jika belum lulus, pengguna menerima hasil kuis dan heart berkurang sesuai aturan aplikasi.")
    h(doc,"3.9 Perancangan Kuis")
    para(doc,"Soal diambil berdasarkan materi dan tingkat lalu diacak. Pilihan ganda serta benar-salah diperiksa secara deterministik. Soal esai mewajibkan proses dan jawaban akhir. Jawaban esai diperiksa melalui normalisasi dan perbandingan numerik bila memungkinkan; jawaban yang membutuhkan evaluasi semantik dikirim ke Gemini berdasarkan rubrik.")
    h(doc,"3.10 Perancangan Gamifikasi")
    para(doc,"Gamifikasi MatHeal mencakup level, poin, streak, heart, misi harian, indikator progres, dan sertifikat. Heart normal memiliki batas lima dan dapat pulih setiap 30 menit ketika berada di bawah batas normal. Misi harian dipilih berdasarkan tanggal dan memberi bonus heart setelah target tercapai. Semua elemen tersebut dihubungkan dengan penyelesaian kuis dan progres belajar.")
    h(doc,"3.11 Perancangan Basis Data")
    table(doc,["Tabel","Fungsi Utama"],[
        ("users","Akun, password hash, role, dan waktu pendaftaran."),("ai_materi","Metadata materi dan lokasi file."),("ai_soal","Pertanyaan, kunci, rubrik, tingkat, tipe, opsi, poin, dan status."),("ai_auth_tokens","Hash token dan masa kedaluwarsa."),("ai_user_progress","Snapshot progres per pengguna."),("ai_aktivitas","Log aktivitas dashboard."),("ai_chat_sessions dan ai_chat_messages","Sesi dan pesan AskMatheal."),("ai_password_reset_requests","Permintaan reset password."),("ai_rate_limits","Pembatasan frekuensi permintaan.")
    ],[2.2,4.8])
    h(doc,"3.12 Implementasi Antarmuka User")
    para(doc,"Antarmuka user dibangun dengan React dan Vite. Landing page menjelaskan manfaat, fitur, alur penggunaan, dan pengembang. Homepage menampilkan materi, progres, streak, heart, misi harian, serta peta level. Halaman materi menampilkan modul PDF dan akses kuis. Halaman kuis menampilkan nomor soal, progres, tipe soal, heart, feedback, dan hasil akhir. Halaman profil menampilkan poin, streak, heart, progres, grafik aktivitas, catatan belajar, serta akses sertifikat. AskMatheal menampilkan percakapan dan riwayat chat.")
    h(doc,"3.13 Implementasi Antarmuka Admin")
    para(doc,"Admin memiliki dashboard, manajemen pengguna, materi, bank soal, video panduan, aktivitas, dan profil. Dashboard menampilkan ringkasan pengguna, aktivitas tujuh hari, materi dibuka, kuis diselesaikan, tren harian, status layanan, dan aktivitas terbaru. Materi dikelola sebagai PDF. Bank soal dikelompokkan berdasarkan materi, tingkat, dan tipe, dengan validasi untuk memastikan opsi soal objektif tersedia serta data esai memiliki struktur yang sesuai.")
    h(doc,"3.14 Implementasi AI Tutor dan Penilaian Esai")
    para(doc,"Gateway AI berbasis Node.js dan Express menerima pertanyaan AskMatheal serta permintaan penilaian esai. Prompt tutor meminta penjelasan matematika dalam bahasa Indonesia dengan langkah yang runtut dan notasi LaTeX. Penilaian esai menerima soal, proses, jawaban akhir, dan kunci referensi. Rubrik terdiri atas ketepatan jawaban akhir 50 poin, validitas metode 40 poin, serta konsistensi proses 10 poin. Jawaban dinyatakan benar bila jawaban akhir benar atau ekuivalen, metode valid, skor minimal 70, dan tidak terdapat kesalahan kritis. Penelitian ini hanya membuktikan perilaku fitur, bukan akurasi penilaiannya terhadap ahli.")
    h(doc,"3.15 Keamanan Implementasi")
    para(doc,"Password disimpan menggunakan password_hash dan diverifikasi menggunakan password_verify. Token disimpan dalam bentuk hash serta memiliki masa kedaluwarsa. Endpoint admin memeriksa role, sedangkan endpoint AI memeriksa token melalui jalur internal. Kredensial server dan GEMINI_KEY disimpan pada environment variable. Sistem menggunakan prepared statement, validasi input, pembatasan unggahan PDF, validasi path, rate limiting, timeout, CORS sesuai konfigurasi, serta header keamanan dasar.")
    h(doc,"3.16 Pelaksanaan Black-Box Testing")
    para(doc,"Black-box testing dilakukan secara mandiri oleh peneliti terhadap 25 skenario fungsi utama. Pengujian membandingkan hasil aktual dengan hasil yang diharapkan tanpa meninjau kode program. Dari dokumen pengujian yang dilampirkan, seluruh 25 skenario berstatus berhasil. Persentase keberhasilan pengujian adalah 25/25 × 100% = 100%. Hasil ini menunjukkan bahwa fungsi yang diuji bekerja sesuai skenario, tetapi tidak membuktikan bahwa aplikasi bebas dari seluruh kemungkinan kesalahan, aman secara menyeluruh, mudah digunakan oleh pengguna, atau akurat secara akademik.")
    cases=[
        ("BB-01","Registrasi","Pendaftaran akun berhasil."),("BB-02","Validasi registrasi","Password kurang dari 8 karakter ditolak."),("BB-03","Login valid","User masuk ke beranda."),("BB-04","Login salah","Sistem menampilkan pesan kesalahan."),("BB-05","Logout","User kembali ke landing page."),("BB-06","Materi","Materi dan PDF terbuka."),("BB-07","Level kuis","Level sedang terkunci sebelum level mudah."),("BB-08","Kuis benar","Feedback benar dan dapat lanjut."),("BB-09","Kuis salah","Feedback salah dan jawaban benar tampil."),("BB-10","Esai tanpa proses","Jawaban tidak dapat dikirim."),("BB-11","Esai valid","Feedback tampil dan lanjut soal."),("BB-12","Kuis lulus","Progres tersimpan dan level terbuka."),("BB-13","Kuis tidak lulus","Heart berkurang dan kuis dapat diulang."),("BB-14","Progres","Progres tetap tersimpan setelah logout."),("BB-15","Misi harian","Status misi selesai dan heart bertambah."),("BB-16","Sertifikat","Opsi sertifikat muncul setelah semua level."),("BB-17","AI Tutor","Pesan gangguan tampil saat layanan menghubungkan kembali."),("BB-18","AI Tutor kosong","Pertanyaan kosong tidak dapat dikirim."),("BB-19","Akses admin user","User dialihkan ke beranda."),("BB-20","Login admin","Admin masuk dashboard."),("BB-21","Materi admin","Materi PDF tersimpan dan tampil."),("BB-22","Unggah non-PDF","Sistem menolak file non-PDF."),("BB-23","Manajemen soal","Soal sesuai tipe dan kunci tersimpan."),("BB-24","Validasi soal","Opsi wajib diisi."),("BB-25","Dashboard","Ringkasan aktivitas dan dashboard tampil.")]
    table(doc,["ID","Fitur","Hasil Aktual Ringkas","Status"],[(a,b,c,"Berhasil") for a,b,c in cases],[.65,1.45,4.15,.75])
    para(doc,"Kolom bukti pada dokumen hasil pengujian asal belum terisi. Sebelum naskah diserahkan, lampirkan screenshot atau respons Postman untuk setiap skenario agar hasil pengujian dapat ditelusuri.")
    h(doc,"3.17 Teknik Analisis Data")
    para(doc,"Data black-box dianalisis secara deskriptif dengan menghitung jumlah skenario berhasil dan gagal. Rumus persentase keberhasilan adalah jumlah skenario berhasil dibagi jumlah seluruh skenario dikali 100 persen. Berdasarkan 25 skenario, tingkat keberhasilan adalah 100 persen. Analisis tidak mencakup skor SUS, korelasi penilaian AI dengan dosen, maupun inferensi statistik karena data tersebut tidak dikumpulkan.")
    h(doc,"3.18 Ringkasan Implementasi Saat Ini")
    para(doc,"MatHeal telah memiliki frontend React, API PHP/MySQL, gateway Node.js untuk Gemini, materi PDF, bank soal, autentikasi berbasis token, progres pengguna, gamifikasi, sertifikat, riwayat AskMatheal, dan dashboard admin. Berdasarkan black-box testing mandiri, 25 fitur atau skenario utama berstatus berhasil. Keterbatasan penelitian adalah tidak adanya evaluasi pengguna dan validator ahli, sehingga kesimpulan penelitian terbatas pada kesesuaian fungsi yang diuji.")


def refs(doc):
    doc.add_page_break(); title(doc,"DAFTAR PUSTAKA PILIHAN")
    items=[
        "Deterding, S., Dixon, D., Khaled, R., & Nacke, L. (2011). From game design elements to gamefulness: Defining gamification. Proceedings of MindTrek, 9-15.",
        "Miao, F., & Holmes, W. (2023). Guidance for generative AI in education and research. UNESCO.",
        "Pressman, R. S., & Maxim, B. R. (2020). Software Engineering: A Practitioner's Approach (9th ed.). McGraw-Hill.",
        "Sailer, M., & Homner, L. (2020). The gamification of learning: A meta-analysis. Educational Psychology Review, 32, 77-112.",
        "ISO/IEC. (2023). ISO/IEC 25010:2023 Systems and software engineering - Product quality model."
    ]
    for x in items:
        p=para(doc,x,False); p.paragraph_format.left_indent=Cm(1); p.paragraph_format.first_line_indent=Cm(-1)


def build():
    doc=Document(); setup(doc); doc.core_properties.title="Revisi Terintegrasi Bab I-III MatHeal"
    chapter1(doc); chapter2(doc); chapter3(doc); refs(doc)
    OUT.parent.mkdir(exist_ok=True); doc.save(OUT); print(OUT)

if __name__=="__main__": build()
