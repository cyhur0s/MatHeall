from pathlib import Path
from docx import Document


SOURCE = Path(r"D:\KULIAH\Penelitian Ilmiah\bab 1.docx")
OUTPUT = Path(r"D:\XAMPP\htdocs\matheal\docs\Bab_1_Terintegrasi_Tanpa_Identifikasi_Rumusan.docx")


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)
    paragraph._p = paragraph._element = None


def main():
    doc = Document(SOURCE)

    # The issue statements in the two removed sections already appear as
    # narrative in the background. Remove only the now-redundant headings,
    # lead-in sentences, and lists to preserve all other existing content.
    remove_texts = {
        "Identifikasi Masalah",
        "Berdasarkan uraian latar belakang, permasalahan yang diidentifikasi dalam penelitian ini adalah:",
        "Platform e-learning yang tersedia bersifat statis dan tidak mampu memberi umpan balik adaptif terhadap kemampuan individual mahasiswa dalam mempelajari matematika informatika.",
        "Tidak tersedianya motivasi berkelanjutan pada platform e-learning yang ada, sehingga mahasiswa cenderung tidak konsisten dalam belajar mandiri di luar jam kuliah.",
        "Penilaian jawaban esai matematika yang masih terbatas pada kecocokan teks dengan kunci jawaban, sehingga berisiko menolak penyelesaian yang sah secara matematis namun berbeda pendekatan.",
        "Ketergantungan fitur AskMatheal dan penilaian jawaban esai pada layanan Google Gemini API berpotensi menimbulkan gangguan ketika API key tidak tersedia atau tidak valid, kuota layanan habis, koneksi internet terganggu, maupun layanan Gemini mengalami kegagalan.",
        "Belum tersedianya aplikasi e-learning matematika informatika berbasis web di Indonesia yang mengintegrasikan gamifikasi multi-elemen dengan AI Tutor berkemampuan penilaian semantik dalam satu sistem yang terintegrasi.",
        "Rumusan Masalah",
        "Berdasarkan identifikasi masalah yang telah diuraikan, rumusan masalah dalam penelitian ini adalah:",
        "Bagaimana merancang dan membangun aplikasi e-learning matematika informatika berbasis web yang mengintegrasikan modul materi terstruktur, kuis bertingkat, gamifikasi, AI Tutor, dan sistem pelacakan progres dalam satu platform?",
        "Bagaimana menerapkan mekanisme gamifikasi berbasis Self-Determination Theory yang dapat mendorong keterlibatan dan konsistensi belajar pengguna?",
        "Bagaimana mengintegrasikan AI Tutor berbasis Google Gemini untuk memberikan penjelasan matematika dan menilai jawaban esai berdasarkan validitas konsep dan proses, bukan sekadar kecocokan teks?",
        "Bagaimana menangani ketergantungan layanan Google Gemini API pada fitur AskMatheal dan penilaian jawaban esai ketika terjadi gangguan layanan atau kegagalan konfigurasi API key?",
        "Bagaimana memastikan seluruh fitur utama aplikasi MatHeal berjalan sesuai spesifikasi melalui pengujian fungsional yang terstruktur?",
    }
    for p in list(doc.paragraphs):
        if p.text.strip() in remove_texts:
            remove_paragraph(p)

    # Align the focus statement with the actual evaluation scope.
    old_focus = "Penelitian ini berfokus pada proses pengembangan dan pengujian kelayakan fungsional aplikasi melalui black box testing terhadap seluruh skenario fitur utama, serta validasi penilaian AI."
    new_focus = (
        "Penelitian ini berfokus pada proses pengembangan dan pengujian kelayakan "
        "fungsional aplikasi melalui black box testing terhadap seluruh skenario fitur utama. "
        "Validasi akademik penilaian AI oleh dosen atau pakar belum dilakukan dan "
        "menjadi rekomendasi untuk penelitian lanjutan."
    )
    focus_para = None
    for p in doc.paragraphs:
        if p.text.strip() == old_focus:
            p.clear()
            p.add_run(new_focus)
            focus_para = p
            break

    # Update the chapter description so it matches the revised structure.
    old_systematic = (
        "BAB I PENDAHULUAN, berisi uraian mengenai latar belakang, identifikasi masalah, "
        "rumusan masalah, batasan masalah, tujuan penelitian, metode penelitian, dan sistematika penulisan."
    )
    new_systematic = (
        "BAB I PENDAHULUAN, berisi uraian mengenai latar belakang, ruang lingkup, "
        "tujuan penelitian, metode penelitian, dan sistematika penulisan."
    )
    for p in doc.paragraphs:
        if p.text.strip() == old_systematic:
            p.clear()
            p.add_run(new_systematic)
            break

    # Correct replacement-character encoding in two existing citations.
    for p in doc.paragraphs:
        if "Ni�o-Rojas" in p.text or "2020�2025" in p.text:
            text = p.text.replace("Ni�o-Rojas", "Niño-Rojas").replace("2020�2025", "2020–2025")
            p.clear()
            p.add_run(text)

    # Leave a concise reviewer note anchored to the edited focus statement.
    if focus_para is not None and hasattr(doc, "add_comment"):
        doc.add_comment(
            focus_para.runs,
            text=(
                "Subbab Identifikasi Masalah dan Rumusan Masalah dilebur ke dalam "
                "Latar Belakang serta Tujuan Penelitian agar struktur Bab I mengikuti pedoman kampus."
            ),
            author="Codex",
            initials="CD",
        )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
