from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = OUT_DIR / "Draf_Penulisan_Ilmiah_MatHeal_Bab_1-4.docx"
ARCH_DIAGRAM = OUT_DIR / "arsitektur_matheal.png"
FLOW_DIAGRAM = OUT_DIR / "alur_belajar_matheal.png"

NAVY = "244A52"
TEAL = "397582"
TEAL_LIGHT = "E8F2F0"
GOLD = "C98950"
GOLD_LIGHT = "F8EFE5"
INK = "1F2F32"
MUTED = "5C6D70"
LINE = "CAD9D6"
PLACEHOLDER = "FFF2CC"
ACADEMIC_TABLE_DXA = 7800


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    # A4 dengan margin kiri 4 cm dan kanan 3 cm menyisakan sekitar 7.937 DXA.
    # Lebar 7.800 DXA + indent 120 DXA menjaga tabel tetap di dalam area teks.
    source_total = sum(widths_dxa)
    scaled = [max(1, round(width * ACADEMIC_TABLE_DXA / source_total)) for width in widths_dxa]
    scaled[-1] += ACADEMIC_TABLE_DXA - sum(scaled)
    widths_dxa = scaled
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Halaman ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def font(size, bold=False):
    try:
        return ImageFont.truetype("arialbd.ttf" if bold else "arial.ttf", size)
    except OSError:
        return ImageFont.load_default()


def draw_box(draw, box, title, subtitle, fill, outline=TEAL):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=24, fill=f"#{fill}", outline=f"#{outline}", width=3)
    title_font = font(30, True)
    subtitle_font = font(21)
    title_box = draw.textbbox((0, 0), title, font=title_font)
    subtitle_box = draw.textbbox((0, 0), subtitle, font=subtitle_font)
    draw.text(((x1+x2-(title_box[2]-title_box[0]))/2, y1+25), title, fill=f"#{INK}", font=title_font)
    draw.text(((x1+x2-(subtitle_box[2]-subtitle_box[0]))/2, y1+72), subtitle, fill=f"#{MUTED}", font=subtitle_font)


def draw_arrow(draw, start, end):
    draw.line([start, end], fill=f"#{TEAL}", width=6)
    ex, ey = end
    draw.polygon([(ex, ey), (ex-16, ey-10), (ex-16, ey+10)], fill=f"#{TEAL}")


def create_diagrams():
    img = Image.new("RGB", (1600, 900), "#FFFEFA")
    d = ImageDraw.Draw(img)
    title = "Arsitektur Sistem MatHeal"
    tb = d.textbbox((0, 0), title, font=font(42, True))
    d.text(((1600-(tb[2]-tb[0]))/2, 42), title, fill=f"#{NAVY}", font=font(42, True))
    draw_box(d, (70, 210, 390, 360), "Pengguna", "Browser responsif", TEAL_LIGHT)
    draw_box(d, (510, 210, 870, 360), "Frontend", "React + Vite + CSS", "F4F8F7")
    draw_box(d, (1010, 130, 1510, 290), "Backend Aplikasi", "PHP API + MySQL", GOLD_LIGHT, GOLD)
    draw_box(d, (1010, 430, 1510, 590), "Layanan AI", "Node.js + Google Gemini", TEAL_LIGHT)
    draw_box(d, (510, 610, 870, 760), "Penyimpanan Lokal", "Heart, streak, cache UI", "F8FAF9")
    draw_arrow(d, (390, 285), (510, 285))
    draw_arrow(d, (870, 260), (1010, 210))
    draw_arrow(d, (870, 310), (1010, 510))
    draw_arrow(d, (690, 360), (690, 610))
    d.text((518, 395), "HTTPS / JSON", fill=f"#{MUTED}", font=font(20, True))
    d.text((1060, 330), "Autentikasi token dan validasi peran", fill=f"#{MUTED}", font=font(20))
    img.save(ARCH_DIAGRAM)

    img2 = Image.new("RGB", (1600, 1180), "#FFFEFA")
    d2 = ImageDraw.Draw(img2)
    title2 = "Alur Belajar Pengguna"
    tb2 = d2.textbbox((0, 0), title2, font=font(42, True))
    d2.text(((1600-(tb2[2]-tb2[0]))/2, 35), title2, fill=f"#{NAVY}", font=font(42, True))
    steps = [
        ("1. Pilih materi", "17 topik matematika informatika"),
        ("2. Pelajari modul", "Teori, rumus, dan contoh soal PDF"),
        ("3. Kerjakan kuis", "10 soal: PG, benar-salah, dan esai"),
        ("4. Terima umpan balik", "Jawaban benar ditampilkan ketika salah"),
        ("5. Evaluasi hasil", "Lulus jika benar minimal 70%"),
        ("6. Lanjutkan level", "Mudah ke sedang ke sulit"),
        ("7. Pantau progres", "Poin, streak, heart, grafik, sertifikat"),
    ]
    y = 130
    for i, (a, b) in enumerate(steps):
        fill = TEAL_LIGHT if i % 2 == 0 else GOLD_LIGHT
        outline = TEAL if i % 2 == 0 else GOLD
        draw_box(d2, (330, y, 1270, y+118), a, b, fill, outline)
        if i < len(steps)-1:
            d2.line([(800, y+118), (800, y+145)], fill=f"#{TEAL}", width=6)
            d2.polygon([(800, y+151), (788, y+133), (812, y+133)], fill=f"#{TEAL}")
        y += 145
    img2.save(FLOW_DIAGRAM)


def set_run_font(run, name="Times New Roman", size=12, bold=None, italic=None, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(4.0)
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(1.25)

    for style_name, size, color, before, after in [
        ("Heading 1", 14, NAVY, 12, 8),
        ("Heading 2", 12, TEAL, 10, 5),
        ("Heading 3", 12, INK, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(12)
        style.paragraph_format.left_indent = Cm(1.25)
        style.paragraph_format.first_line_indent = Cm(-0.63)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.5

    if "Caption Academic" not in styles:
        cap = styles.add_style("Caption Academic", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = "Times New Roman"
        cap._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        cap._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        cap.font.size = Pt(10)
        cap.font.bold = True
        cap.font.color.rgb = RGBColor.from_string(MUTED)
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(6)
        cap.paragraph_format.space_after = Pt(6)
        cap.paragraph_format.first_line_indent = Cm(0)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "DRAF PENULISAN ILMIAH  |  MATHEAL"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(hp.runs[0], size=8.5, bold=True, color=MUTED)
    footer = section.footer
    add_page_number(footer.paragraphs[0])
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""


def add_body(doc, text, bold_prefix=None, italic=False, no_indent=False):
    p = doc.add_paragraph()
    if no_indent:
        p.paragraph_format.first_line_indent = Cm(0)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.first_line_indent = Cm(-0.63)
        set_run_font(p.add_run(item))


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.first_line_indent = Cm(-0.63)
        set_run_font(p.add_run(item))


def add_placeholder(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PLACEHOLDER)
    p_pr.append(shd)
    r = p.add_run("BAGIAN YANG HARUS DIISI PENELITI: " + text)
    set_run_font(r, size=10.5, bold=True, color="7A5A00")
    return p


def add_chapter(doc, title):
    doc.add_page_break()
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(title)
    set_run_font(r, size=14, bold=True, color=NAVY)
    return p


def add_section(doc, title, level=2):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(title)
    set_run_font(r, size=12 if level > 1 else 14, bold=True, color=TEAL if level == 2 else INK)
    return p


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for i, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], TEAL)
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=10, bold=True, color="FFFFFF")
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if ridx % 2 == 1:
                set_cell_shading(cells[i], "F7FAF9")
            p = cells[i].paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=9.5)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc, path, caption, width_cm=13.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Cm(width_cm))
    shape._inline.docPr.set("title", caption)
    shape._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph(style="Caption Academic")
    cap.add_run(caption)


def add_cover(doc):
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("DRAF PENULISAN ILMIAH")
    set_run_font(r, size=15, bold=True, color=TEAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("PENGEMBANGAN APLIKASI E-LEARNING MATEMATIKA INFORMATIKA BERBASIS WEB \"MATHEAL\" DENGAN AI TUTOR DAN GAMIFIKASI")
    set_run_font(r, size=16, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("BAB I sampai BAB IV")
    set_run_font(r, size=12, bold=True, color=GOLD)

    # Monogram sederhana sebagai pengganti logo kampus yang belum diberikan.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("M")
    set_run_font(r, name="Arial", size=38, bold=True, color=TEAL)

    for line in [
        "Disusun oleh:",
        "[NAMA LENGKAP]",
        "[NPM]",
        "",
        "PROGRAM STUDI [NAMA PROGRAM STUDI]",
        "FAKULTAS [NAMA FAKULTAS]",
        "[NAMA UNIVERSITAS]",
        "2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(line)
        set_run_font(r, size=12, bold=line.startswith("[") or line.startswith("PROGRAM") or line.startswith("FAKULTAS"))


def add_front_matter(doc):
    doc.add_page_break()
    add_section(doc, "CATATAN PENGGUNAAN DRAF", 1)
    add_body(doc, "Dokumen ini disusun berdasarkan implementasi aplikasi MatHeal pada folder proyek. Draf telah memuat narasi akademik, rancangan pengujian, dan referensi awal, tetapi belum boleh diperlakukan sebagai laporan hasil penelitian final sebelum data penelitian benar-benar dikumpulkan.")
    add_body(doc, "Judul yang direkomendasikan lebih spesifik daripada judul awal karena menyebutkan ruang lingkup matematika informatika dan pendekatan gamifikasi. Apabila pedoman kampus membatasi panjang judul, judul awal \"Pengembangan Aplikasi E-Learning Matematika Berbasis Web MatHeal dengan AI Tutor\" tetap dapat digunakan.")
    add_placeholder(doc, "Lengkapi identitas penulis, institusi, dosen pembimbing, lokasi penelitian, waktu penelitian, jumlah responden, hasil black-box testing, hasil SUS, dan bukti tangkapan layar aplikasi.")
    add_body(doc, "Saran metodologis utama adalah menggunakan metode Research and Development dengan model prototipe. Pengujian produk dilakukan melalui black-box testing, pengujian responsif, evaluasi kualitas terpilih dari ISO/IEC 25010:2023, serta System Usability Scale (SUS). Jika penelitian ingin menyimpulkan peningkatan hasil belajar, wajib ditambahkan desain pretest-posttest; tanpa itu, kesimpulan hanya boleh membahas keberhasilan pengembangan dan kelayakan penggunaan.")

    add_section(doc, "DAFTAR ISI RINGKAS", 1)
    toc_rows = [
        ("BAB I", "Pendahuluan"),
        ("BAB II", "Tinjauan Pustaka"),
        ("BAB III", "Analisis, Perancangan, Implementasi, dan Pengujian"),
        ("BAB IV", "Penutup: Kesimpulan dan Saran"),
        ("-", "Daftar Pustaka"),
    ]
    add_table(doc, ["Bagian", "Isi"], toc_rows, [1500, 7860])


def chapter_one(doc):
    add_chapter(doc, "BAB I\nPENDAHULUAN")
    add_section(doc, "1.1 Latar Belakang")
    add_body(doc, "Perkembangan teknologi digital telah memperluas cara mahasiswa memperoleh materi, berlatih, dan menerima umpan balik. Pembelajaran berbasis web memungkinkan sumber belajar diakses tanpa terikat tempat serta dapat mengintegrasikan teks, rumus, kuis, dan rekaman progres dalam satu lingkungan. Dalam konteks pendidikan tinggi, pendekatan ini relevan untuk mendukung pembelajaran mandiri, terutama pada mata kuliah yang membutuhkan latihan berulang dan pemahaman proses seperti matematika informatika.")
    add_body(doc, "Matematika informatika memuat konsep yang menjadi dasar penalaran komputasional, antara lain logika matematika, himpunan, fungsi, bilangan biner, matriks, graf, relasi rekurensi, dan analisis algoritma. Bagi pengguna pemula, kesulitan tidak hanya terletak pada mengingat rumus, tetapi juga pada memilih konsep, menyusun langkah penyelesaian, dan memeriksa kebenaran hasil. Modul statis dapat menjelaskan teori, namun sering belum memberikan umpan balik langsung ketika pengguna melakukan kesalahan atau menggunakan metode yang berbeda dari contoh.")
    add_body(doc, "Sistem tutor cerdas dapat melengkapi pembelajaran dengan umpan balik yang lebih personal. Kajian Niño-Rojas dkk. (2024) terhadap 43 publikasi menunjukkan bahwa intelligent tutoring system dalam matematika digunakan untuk mendukung pemahaman konsep dan pengembangan keterampilan melalui lingkungan adaptif. Kajian yang lebih baru juga menekankan bahwa peran tutor digital perlu tetap terhubung dengan tujuan pedagogis, bukan sekadar menggantikan penyajian soal. Karena itu, AI pada aplikasi pendidikan perlu diposisikan sebagai pendamping belajar yang membantu menjelaskan langkah dan menilai alasan, bukan sebagai pengganti dosen.")
    add_body(doc, "Selain dukungan tutor, konsistensi belajar merupakan persoalan penting. Gamifikasi memindahkan elemen desain permainan ke konteks nonpermainan (Deterding dkk., 2011). Meta-analisis Sailer dan Homner (2020) menemukan efek positif berukuran kecil hingga sedang terhadap hasil kognitif, motivasional, dan perilaku, sedangkan Bai, Hew, dan Huang (2020) melaporkan bahwa dampaknya bergantung pada desain dan konteks. Temuan tersebut mendukung penggunaan level, progres, streak, heart, misi, dan sertifikat secara terarah, dengan catatan bahwa elemen tersebut harus memperkuat aktivitas belajar, bukan menjadi tujuan utama.")
    add_body(doc, "Di sisi lain, penggunaan generative AI membawa risiko kesalahan informasi, ketergantungan, privasi, dan keputusan penilaian yang tidak transparan. UNESCO (Miao & Holmes, 2023) merekomendasikan pendekatan yang berpusat pada manusia, aman, etis, dan bermakna. Oleh sebab itu, integrasi AI pada MatHeal dirancang dengan pembatasan fungsi: AskMatheal memberi penjelasan matematika berbahasa Indonesia, sedangkan penilai esai memeriksa jawaban akhir dan validitas metode menggunakan rubrik terstruktur. Kunci jawaban diperlakukan sebagai referensi, sehingga metode alternatif yang sah dapat diterima.")
    add_body(doc, "Berdasarkan kebutuhan tersebut, dikembangkan MatHeal, yaitu aplikasi e-learning matematika informatika berbasis web dengan AI Tutor dan gamifikasi. Aplikasi menyediakan 17 materi, masing-masing dengan tingkat mudah, sedang, dan sulit. Setiap sesi kuis berisi 10 soal acak yang menggabungkan pilihan ganda, benar-salah, dan esai. Pengguna dapat membaca modul, mengerjakan kuis, memperoleh jawaban benar ketika salah, membuka level berikutnya setelah memenuhi batas kelulusan 70%, memantau progres, serta mencetak sertifikat setelah menuntaskan seluruh level suatu materi. Administrator dapat mengelola pengguna, materi, bank soal, dan aktivitas melalui dashboard.")
    add_body(doc, "Pengembangan ini diharapkan menghasilkan lingkungan belajar yang lebih terstruktur bagi pengguna awam. Penelitian tidak hanya membahas implementasi teknis, tetapi juga menilai kesesuaian fungsi, kualitas antarmuka, dan penerimaan pengguna. Dengan demikian, judul yang diusulkan adalah \"Pengembangan Aplikasi E-Learning Matematika Informatika Berbasis Web MatHeal dengan AI Tutor dan Gamifikasi\".")

    add_section(doc, "1.2 Identifikasi Masalah")
    add_bullets(doc, [
        "Pengguna pemula membutuhkan materi matematika informatika yang tersusun bertahap, memuat teori, rumus, dan contoh soal.",
        "Latihan yang hanya menggunakan satu bentuk soal kurang variatif dan kurang melatih pemahaman dari berbagai sudut.",
        "Pemeriksaan esai berbasis kecocokan teks tidak mampu menerima cara alternatif yang sah secara matematika.",
        "Pengguna membutuhkan umpan balik segera, rekaman progres, serta mekanisme yang mendorong konsistensi belajar.",
        "Administrator membutuhkan pengelolaan materi dan soal yang selaras dengan aktivitas pengguna serta ringkasan data yang mudah dipantau.",
        "Aplikasi perlu tetap nyaman digunakan pada layar desktop, laptop, tablet, dan telepon seluler.",
    ])

    add_section(doc, "1.3 Rumusan Masalah")
    add_numbers(doc, [
        "Bagaimana merancang dan membangun aplikasi e-learning matematika informatika berbasis web yang mengintegrasikan materi, kuis bertingkat, progres, dan administrasi konten?",
        "Bagaimana mengintegrasikan Google Gemini sebagai AI Tutor dan penilai jawaban esai yang mempertimbangkan jawaban akhir serta validitas metode alternatif?",
        "Bagaimana menerapkan gamifikasi agar mendukung urutan belajar mudah-sedang-sulit dan menjaga motivasi tanpa mengaburkan tujuan pembelajaran?",
        "Bagaimana menguji fungsionalitas, responsivitas, dan usability aplikasi MatHeal sebelum digunakan oleh lebih banyak pengguna?",
    ])

    add_section(doc, "1.4 Batasan Masalah")
    add_bullets(doc, [
        "Materi berfokus pada 17 topik matematika dan matematika informatika yang tersedia pada katalog MatHeal.",
        "Setiap materi memiliki tiga tingkat, yaitu mudah, sedang, dan sulit; level berikutnya terbuka setelah level sebelumnya dinyatakan lulus.",
        "Setiap kuis terdiri atas 10 soal acak dengan mode pilihan ganda, benar-salah, dan esai yang diambil dari bank soal administrator.",
        "Nilai minimum kelulusan adalah 70% jawaban benar.",
        "AI yang digunakan adalah Google Gemini melalui layanan backend Node.js. AI tidak menggantikan keputusan akademik dosen dan hasilnya harus diperlakukan sebagai umpan balik otomatis.",
        "Aplikasi dikembangkan dengan React pada frontend, PHP dan MySQL pada layanan aplikasi, serta Node.js/Express pada layanan AI.",
        "Penelitian tahap ini menilai keberhasilan pengembangan dan usability. Klaim peningkatan hasil belajar hanya dapat dibuat jika peneliti menambahkan eksperimen pretest-posttest.",
    ])

    add_section(doc, "1.5 Tujuan Penelitian")
    add_numbers(doc, [
        "Mengembangkan aplikasi e-learning matematika informatika berbasis web yang menghubungkan materi, kuis, progres, sertifikat, dan dashboard admin.",
        "Menerapkan AI Tutor untuk memberi penjelasan matematika serta menilai jawaban esai berdasarkan kebenaran hasil dan validitas proses.",
        "Menerapkan gamifikasi berupa level, poin, streak, heart, misi harian, progres, dan sertifikat untuk mendukung keteraturan belajar.",
        "Mengevaluasi fungsi utama, kompatibilitas tampilan, dan usability MatHeal menggunakan skenario pengujian yang terukur.",
    ])

    add_section(doc, "1.6 Manfaat Penelitian")
    add_section(doc, "1.6.1 Manfaat bagi Pengguna", 3)
    add_body(doc, "MatHeal memberi jalur belajar yang jelas dari materi ke kuis, umpan balik langsung, serta informasi perkembangan belajar. Pengguna pemula dapat memulai dari level mudah dan mengulang kuis tanpa kehilangan akses setelah lulus.")
    add_section(doc, "1.6.2 Manfaat bagi Pengajar atau Administrator", 3)
    add_body(doc, "Administrator dapat memelihara materi dan bank soal, memantau aktivitas pengguna, serta menjaga distribusi tingkat dan tipe soal. Data aktivitas dapat digunakan sebagai bahan evaluasi konten tanpa menyimpan rincian jawaban pribadi yang tidak diperlukan.")
    add_section(doc, "1.6.3 Manfaat Akademik dan Pengembangan", 3)
    add_body(doc, "Penelitian ini dapat menjadi referensi implementasi AI generatif pada e-learning matematika dengan mekanisme penilaian esai berbasis rubrik dan pengamanan prompt. Arsitektur modularnya juga membuka peluang pengembangan adaptivitas, analitik pembelajaran, dan eksperimen efektivitas belajar pada penelitian lanjutan.")

    add_section(doc, "1.7 Metode Penelitian Secara Ringkas")
    add_body(doc, "Penelitian menggunakan pendekatan Research and Development dengan model prototipe. Tahapannya meliputi identifikasi kebutuhan, pembuatan prototipe, evaluasi prototipe, penyempurnaan, implementasi, dan pengujian. Data kebutuhan diperoleh melalui studi pustaka, observasi proses belajar, dan analisis aplikasi. Evaluasi produk direncanakan menggunakan black-box testing, pengujian responsif, indikator kualitas ISO/IEC 25010:2023 yang relevan, serta System Usability Scale.")
    add_placeholder(doc, "Tuliskan subjek penelitian, teknik sampling, jumlah responden, lokasi, dan rentang waktu pengambilan data sesuai pelaksanaan sebenarnya.")

    add_section(doc, "1.8 Sistematika Penulisan")
    add_body(doc, "BAB I membahas latar belakang, permasalahan, batasan, tujuan, manfaat, metode singkat, dan sistematika. BAB II menguraikan teori e-learning, matematika informatika, gamifikasi, AI Tutor, penilaian esai, teknologi web, kualitas perangkat lunak, serta penelitian terkait. BAB III menjelaskan metode pengembangan, kebutuhan, arsitektur, basis data, implementasi fungsi, dan rancangan pengujian. BAB IV memuat kesimpulan, keterbatasan, dan saran pengembangan.")


def chapter_two(doc):
    add_chapter(doc, "BAB II\nTINJAUAN PUSTAKA")
    add_section(doc, "2.1 E-Learning Berbasis Web")
    add_body(doc, "E-learning merupakan proses pembelajaran yang memanfaatkan teknologi digital untuk menyampaikan materi, aktivitas, komunikasi, dan evaluasi. Pada aplikasi berbasis web, pengguna cukup menggunakan browser sehingga distribusi sistem lebih mudah dibanding aplikasi yang harus dipasang pada setiap perangkat. Nilai utama e-learning bukan hanya digitalisasi bahan ajar, tetapi integrasi antara tujuan, konten, aktivitas, umpan balik, dan pemantauan kemajuan.")
    add_body(doc, "Dalam MatHeal, prinsip tersebut diwujudkan melalui hubungan langsung antara modul PDF, pemilihan materi, kuis bertingkat, umpan balik per soal, hasil kuis, profil progres, dan sertifikat. Materi bukan lampiran terpisah, melainkan tahap belajar sebelum latihan. Kuis juga tidak berdiri sendiri karena tingkat berikutnya dipengaruhi oleh status kelulusan tingkat sebelumnya.")

    add_section(doc, "2.2 Matematika Informatika")
    add_body(doc, "Matematika informatika adalah kumpulan konsep matematika yang mendukung representasi data, penalaran logis, pemodelan, dan analisis komputasi. Ruang lingkup pada MatHeal meliputi limit, turunan, himpunan dan fungsi komposisi, logika proposisi, fungsi Boolean dan peta Karnaugh, bilangan kompleks, matriks, transformasi linier, bilangan biner, relasi rekurensi, graf, analisis algoritma, integral, persamaan linear, geometri, logika matematika, dan trigonometri.")
    add_body(doc, "Pembelajaran topik tersebut membutuhkan representasi simbolis dan latihan prosedural, tetapi juga penjelasan konseptual. Karena itu, materi idealnya memuat prasyarat, definisi, rumus umum, contoh bertahap, kesalahan umum, dan latihan. Kuis perlu menilai pengenalan konsep melalui pilihan ganda, validasi pernyataan melalui benar-salah, serta proses penalaran melalui esai.")

    add_section(doc, "2.3 Gamifikasi")
    add_body(doc, "Deterding dkk. (2011) mendefinisikan gamifikasi sebagai penggunaan elemen desain permainan dalam konteks nonpermainan. Pada pendidikan, elemen yang umum meliputi poin, level, tantangan, progres, penghargaan, dan batas sumber daya. Elemen tersebut bekerja sebagai informasi tujuan dan umpan balik perilaku; efektivitasnya bergantung pada hubungan antara mekanik permainan dan aktivitas belajar.")
    add_body(doc, "Sailer dan Homner (2020) menemukan efek positif gamifikasi terhadap hasil kognitif (g = 0,49), motivasional (g = 0,36), dan perilaku (g = 0,25). Bai dkk. (2020) juga melaporkan peningkatan performa rata-rata, tetapi menekankan bahwa pengalaman pengguna dapat positif maupun negatif. Oleh sebab itu, desain MatHeal menghindari leaderboard kompetitif dan menekankan kemajuan pribadi.")
    add_table(doc, ["Elemen", "Penerapan di MatHeal", "Tujuan Pembelajaran"], [
        ("Level", "Mudah, sedang, sulit", "Menyusun tahapan penguasaan"),
        ("Poin", "100 poin untuk setiap level lulus", "Menampilkan pencapaian kumulatif"),
        ("Streak", "Bertambah jika kuis diselesaikan pada hari berurutan WIB", "Mendorong konsistensi"),
        ("Heart", "Berkurang saat tidak lulus dan pulih setiap 30 menit", "Memberi konsekuensi ringan dan jeda refleksi"),
        ("Misi harian", "Target harian berganti dan memberi bonus heart", "Mendorong aktivitas yang terarah"),
        ("Sertifikat", "Tersedia setelah tiga level suatu materi selesai", "Memberi pengakuan penyelesaian"),
    ], [1550, 3600, 4210])

    add_section(doc, "2.4 Intelligent Tutoring System dan AI Tutor")
    add_body(doc, "Intelligent tutoring system (ITS) dirancang untuk mendukung pembelajaran individual melalui penyajian aktivitas dan umpan balik. Niño-Rojas dkk. (2024) menunjukkan pertumbuhan penggunaan ITS pada pendidikan matematika dan menyoroti tujuan, keuntungan, efektivitas, keterbatasan, serta metodologi pengembangannya. Kajian ITS matematika periode 2003-2023 juga menunjukkan bahwa banyak sistem masih berada pada tingkat augmentasi, yaitu meningkatkan fungsi pembelajaran yang ada, belum sepenuhnya mendefinisikan ulang proses pembelajaran.")
    add_body(doc, "AskMatheal termasuk dukungan augmentatif. Pengguna dapat meminta penjelasan, langkah penyelesaian, dan klarifikasi konsep dalam bahasa Indonesia. Prompt sistem mengarahkan AI untuk menggunakan format yang mudah dibaca dan notasi LaTeX. Ketika layanan AI tidak tersedia, aplikasi memberikan pesan yang transparan dan menyarankan langkah belajar umum, bukan berpura-pura menghasilkan jawaban AI.")
    add_body(doc, "Google Gemini diakses melalui metode generateContent pada Gemini API. Dokumentasi Google menjelaskan bahwa permintaan terdiri atas konten yang dikirim ke model dan respons kandidat yang dihasilkan. Pada MatHeal, kunci API disimpan pada server Node.js, bukan pada frontend, dan endpoint AI memerlukan token pengguna yang diverifikasi melalui backend PHP.")

    add_section(doc, "2.5 Penilaian Jawaban Esai Berbantuan AI")
    add_body(doc, "Jawaban matematika dapat dinyatakan dengan notasi, urutan, atau strategi berbeda. Pencocokan teks secara langsung berisiko menolak jawaban ekuivalen. MatHeal menggunakan dua lapis pemeriksaan. Pertama, pemeriksaan deterministik menerima kecocokan teks yang telah dinormalisasi dan nilai numerik yang sama dalam toleransi 0,001. Kedua, jika jawaban memerlukan penilaian semantik, Gemini menilai dengan rubrik terstruktur.")
    add_table(doc, ["Komponen Rubrik", "Bobot", "Kriteria"], [
        ("Jawaban akhir", "50", "Benar atau ekuivalen secara matematis dengan referensi"),
        ("Metode", "40", "Konsep, rumus, dan langkah utama valid"),
        ("Konsistensi", "10", "Proses konsisten dengan jawaban akhir"),
    ], [2500, 1100, 5760])
    add_body(doc, "Jawaban dinyatakan benar apabila jawaban akhir benar, metode valid, skor minimal 70, dan tidak terdapat kesalahan kritis. Prompt juga menandai isi soal, proses, jawaban, dan kunci sebagai data yang tidak boleh diikuti sebagai instruksi. Strategi ini mengurangi risiko prompt injection, tetapi tidak menghilangkan kemungkinan kesalahan model. Karena itu, bank soal tetap memerlukan kunci dan rubrik yang berkualitas, serta evaluasi sampel jawaban oleh dosen.")

    add_section(doc, "2.6 Umpan Balik Formatif")
    add_body(doc, "Umpan balik formatif diberikan selama proses belajar agar pengguna mengetahui kesenjangan antara respons dan jawaban yang diharapkan. Pada MatHeal, jawaban salah tidak menghentikan sesi dan tidak lagi hanya memberikan hint. Sistem menampilkan jawaban yang benar dan feedback singkat, kemudian pengguna melanjutkan ke soal berikutnya. Hasil akhir mencatat jumlah benar dan salah. Pendekatan ini menjaga alur 10 soal tetap selesai dan memberi kesempatan melakukan evaluasi setelah satu sesi utuh.")

    add_section(doc, "2.7 Teknologi Pengembangan")
    add_table(doc, ["Teknologi", "Peran pada Sistem"], [
        ("React", "Membangun antarmuka berbasis komponen, state, navigasi, dan interaksi pengguna"),
        ("Vite", "Menjalankan lingkungan pengembangan dan membangun aset produksi frontend"),
        ("CSS", "Menerapkan desain responsif, tipografi, kartu, dan tampilan gamifikasi"),
        ("PHP", "Menyediakan endpoint autentikasi, CRUD, aktivitas, progres, dan akses materi"),
        ("MySQL/MariaDB", "Menyimpan akun, materi, soal, token, aktivitas, progres, dan percakapan"),
        ("Node.js/Express", "Menyediakan layanan AskMatheal dan penilaian esai"),
        ("Google Gemini", "Menghasilkan penjelasan tutor dan evaluasi semantik jawaban esai"),
        ("Chart.js", "Menampilkan grafik progres pengguna"),
    ], [2300, 7060])
    add_body(doc, "Dokumentasi React menjelaskan bahwa aplikasi dibangun dari komponen yang memiliki logika dan tampilan sendiri serta dapat mengelola state untuk merespons interaksi. Pola tersebut sesuai dengan struktur MatHeal yang memisahkan landing page, login, homepage, materi, kuis, profil, AskMatheal, dan halaman admin.")

    add_section(doc, "2.8 Keamanan dan Privasi Dasar")
    add_body(doc, "Keamanan aplikasi pendidikan perlu melindungi akun, kunci AI, dan data aktivitas. MatHeal menggunakan password hash, prepared statement pada operasi basis data, token sesi yang memiliki masa berlaku, verifikasi role, pembatasan origin, rate limiting, serta pemisahan kunci Gemini pada server. Dokumentasi PHP menjelaskan bahwa prepared statement memisahkan template perintah dari nilai parameter dan membantu melindungi aplikasi dari SQL injection.")
    add_body(doc, "Data aktivitas admin dibatasi pada kejadian seperti pendaftaran, login, logout, pembukaan materi, dan penyelesaian kuis. Implementasi saat ini tidak dimaksudkan untuk menyimpan rincian jawaban atau nilai setiap percobaan. Prinsip minimisasi data ini perlu dipertahankan pada deployment dan dijelaskan dalam kebijakan privasi.")

    add_section(doc, "2.9 Kualitas Perangkat Lunak")
    add_body(doc, "ISO/IEC 25010:2023 mendefinisikan model kualitas produk perangkat lunak dengan sembilan karakteristik yang dapat digunakan untuk menetapkan kebutuhan, tujuan pengujian, kontrol kualitas, dan kriteria penerimaan. Penelitian ini memilih karakteristik yang relevan dengan ruang lingkup: functional suitability, usability, reliability, security, performance efficiency, dan compatibility. Pemilihan sebagian karakteristik harus dijelaskan agar evaluasi tetap realistis dan terukur.")

    add_section(doc, "2.10 System Usability Scale")
    add_body(doc, "System Usability Scale (SUS) dikembangkan oleh Brooke (1996) sebagai instrumen 10 pernyataan dengan skala Likert lima poin yang menghasilkan skor 0-100. Untuk item ganjil, kontribusi skor adalah jawaban dikurangi satu; untuk item genap, kontribusi skor adalah lima dikurangi jawaban. Jumlah kontribusi dikalikan 2,5. Skor SUS bukan persentase keberhasilan tugas, tetapi ukuran persepsi usability secara keseluruhan.")

    add_section(doc, "2.11 Penelitian Terkait")
    add_table(doc, ["Penelitian", "Temuan Utama", "Relevansi"], [
        ("Sailer & Homner (2020)", "Gamifikasi memberi efek positif pada aspek kognitif, motivasional, dan perilaku", "Landasan level, progres, dan penghargaan"),
        ("Bai, Hew, & Huang (2020)", "Gamifikasi meningkatkan performa rata-rata, tetapi dampaknya dipengaruhi desain", "Alasan mengutamakan kemajuan pribadi"),
        ("Niño-Rojas dkk. (2024)", "ITS matematika mendukung lingkungan adaptif dan pemahaman konsep", "Landasan AI Tutor matematika"),
        ("Miao & Holmes (2023)", "GenAI pendidikan harus human-centered, aman, etis, dan bermakna", "Landasan pembatasan AI dan transparansi"),
        ("ISO/IEC 25010:2023", "Model kualitas produk untuk menetapkan dan mengevaluasi kualitas perangkat lunak", "Landasan pengujian kualitas MatHeal"),
    ], [2450, 3800, 3110])

    add_section(doc, "2.12 Kerangka Berpikir")
    add_body(doc, "Masalah awal berupa materi yang tersebar, latihan kurang variatif, keterbatasan umpan balik, dan sulitnya memantau kemajuan. Solusi yang dirancang adalah platform web yang menggabungkan modul, kuis multimode dan bertingkat, AI Tutor, penilaian esai berbasis rubrik, serta gamifikasi. Proses pengembangan menggunakan model prototipe. Produk kemudian dinilai secara fungsional dan melalui persepsi pengguna. Keluaran penelitian berupa aplikasi MatHeal dan informasi kelayakan awalnya sebagai media belajar pendamping.")


def chapter_three(doc):
    add_chapter(doc, "BAB III\nANALISIS, PERANCANGAN, IMPLEMENTASI, DAN PENGUJIAN")
    add_section(doc, "3.1 Jenis dan Pendekatan Penelitian")
    add_body(doc, "Penelitian ini merupakan Research and Development yang bertujuan menghasilkan dan mengevaluasi produk perangkat lunak pendidikan. Model pengembangan yang digunakan adalah prototipe karena kebutuhan antarmuka dan alur belajar mengalami penyempurnaan berulang berdasarkan evaluasi. Tahapannya terdiri atas komunikasi kebutuhan, perancangan cepat, pembangunan prototipe, evaluasi pengguna, revisi, implementasi, dan pengujian.")

    add_section(doc, "3.2 Objek, Subjek, dan Lokasi Penelitian")
    add_body(doc, "Objek penelitian adalah aplikasi MatHeal. Subjek uji yang disarankan terdiri atas mahasiswa yang sedang atau pernah mempelajari matematika informatika, dosen/pengajar sebagai validator materi dan soal, serta administrator sebagai pengguna dashboard. Pengujian dapat dilakukan di laboratorium komputer atau secara daring dengan perangkat responden sendiri.")
    add_placeholder(doc, "Isi nama kampus/lokasi, populasi, teknik sampling, kriteria responden, jumlah responden, serta tanggal pelaksanaan. Jangan menuliskan jumlah sampel yang belum benar-benar direkrut.")

    add_section(doc, "3.3 Teknik Pengumpulan Data")
    add_numbers(doc, [
        "Studi pustaka terhadap e-learning, gamifikasi, intelligent tutoring system, AI generatif, usability, dan kualitas perangkat lunak.",
        "Observasi terhadap alur belajar pengguna dan pengelolaan konten oleh admin.",
        "Dokumentasi berupa struktur kode, basis data, modul materi, bank soal, dan hasil pengujian.",
        "Kuesioner SUS untuk mengukur persepsi usability setelah responden menyelesaikan skenario tugas.",
        "Validasi ahli untuk menilai kesesuaian materi, soal, kunci, dan rubrik penilaian esai.",
    ])

    add_section(doc, "3.4 Tahapan Pengembangan Prototipe")
    add_table(doc, ["Tahap", "Aktivitas", "Keluaran"], [
        ("Komunikasi", "Mengidentifikasi kebutuhan user dan admin", "Daftar masalah dan kebutuhan"),
        ("Perancangan cepat", "Menyusun alur, struktur halaman, data, dan aturan gamifikasi", "Wireframe dan model data"),
        ("Pembangunan prototipe", "Mengimplementasikan frontend, PHP API, basis data, dan layanan AI", "Prototipe fungsional"),
        ("Evaluasi", "Memeriksa navigasi, soal, feedback, progres, dan tampilan responsif", "Daftar temuan"),
        ("Penyempurnaan", "Memperbaiki fungsi, CSS, sinkronisasi, dan keamanan", "Versi kandidat rilis"),
        ("Pengujian", "Black-box, responsif, usability, dan validasi konten", "Data evaluasi produk"),
    ], [1600, 4750, 3010])

    add_section(doc, "3.5 Analisis Kebutuhan Pengguna")
    add_section(doc, "3.5.1 Kebutuhan Fungsional User", 3)
    add_bullets(doc, [
        "Mendaftar, login, meminta reset password, dan logout.",
        "Memilih satu dari 17 materi serta membaca modul PDF yang sesuai.",
        "Memilih tingkat yang telah terbuka dan mengerjakan 10 soal acak dalam tiga mode.",
        "Menerima jawaban benar dan feedback ketika jawaban salah.",
        "Menerima penilaian esai yang mempertimbangkan proses dan metode alternatif.",
        "Mengulang kuis, membuka level berikutnya, dan tetap dapat mengerjakan ulang level yang telah lulus.",
        "Melihat poin, streak, rekor streak, heart, misi, progres, grafik aktivitas, dan catatan belajar.",
        "Mengakses AskMatheal untuk meminta penjelasan dan melihat riwayat percakapan.",
        "Mengisi nama lengkap dan tanggal lalu mencetak sertifikat per materi yang selesai.",
    ])
    add_section(doc, "3.5.2 Kebutuhan Fungsional Admin", 3)
    add_bullets(doc, [
        "Melihat ringkasan pengguna, aktivitas, materi, kuis, tren tujuh hari, dan status layanan.",
        "Mengelola akun pengguna dan permintaan reset password.",
        "Mengelola 17 materi dan berkas PDF.",
        "Mengelola bank soal berdasarkan materi, tingkat, dan tipe soal.",
        "Mengelola video panduan dan informasi developer.",
        "Melihat aktivitas user seperti registrasi, login, logout, membuka materi, dan menyelesaikan kuis.",
    ])
    add_section(doc, "3.5.3 Kebutuhan Nonfungsional", 3)
    add_table(doc, ["Aspek", "Kebutuhan"], [
        ("Usability", "Navigasi mudah dipahami pengguna awam; istilah dan feedback konsisten"),
        ("Compatibility", "Tampilan menyesuaikan desktop, laptop, tablet, dan mobile"),
        ("Security", "Password di-hash; token dan role diverifikasi; kunci AI tidak berada di frontend"),
        ("Reliability", "Kesalahan layanan AI ditangani tanpa merusak sesi kuis"),
        ("Performance", "Aset dibangun untuk produksi; API memiliki timeout dan koneksi persisten"),
        ("Maintainability", "Frontend, PHP API, server AI, data katalog, dan utilitas dipisahkan"),
    ], [2200, 7160])

    add_section(doc, "3.6 Arsitektur Sistem")
    add_figure(doc, ARCH_DIAGRAM, "Gambar 3.1 Arsitektur sistem MatHeal")
    add_body(doc, "Frontend React menampilkan antarmuka dan mengelola interaksi. Permintaan data akun, materi, soal, aktivitas, serta progres dikirim ke backend PHP dan disimpan pada MySQL. Pertanyaan AskMatheal dan jawaban esai dikirim ke server Node.js/Express, lalu server memanggil Google Gemini. Sebelum menggunakan endpoint AI, token pengguna diverifikasi melalui backend PHP. Sebagian state gamifikasi disimpan pada localStorage dan disinkronkan ke progres akun.")

    add_section(doc, "3.7 Perancangan Alur Belajar")
    add_figure(doc, FLOW_DIAGRAM, "Gambar 3.2 Alur belajar pengguna MatHeal", width_cm=13.2)
    add_body(doc, "Pengguna memulai dari pemilihan materi. Halaman materi menyediakan akses modul dan tombol untuk memulai kuis terkait. Tingkat sedang hanya terbuka setelah mudah lulus, sedangkan tingkat sulit terbuka setelah sedang lulus. Setelah 10 soal selesai, sistem menghitung jumlah benar. Jika benar minimal tujuh, level lulus; jika belum, heart berkurang satu dan pengguna dapat mengulang ketika masih memiliki heart.")

    add_section(doc, "3.8 Perancangan Kuis")
    add_body(doc, "Endpoint kuis mengambil soal berdasarkan materi dan tingkat lalu mengacak hasil. Paket 10 soal disusun agar tiga mode tetap terwakili secara seimbang. Pilihan ganda dan benar-salah dinilai deterministik. Esai mewajibkan proses dan jawaban akhir, kemudian diperiksa secara deterministik atau dikirim ke Gemini untuk evaluasi semantik.")
    add_table(doc, ["Kondisi", "Respons Sistem"], [
        ("Jawaban objektif benar", "Menampilkan status benar dan lanjut ke soal berikutnya"),
        ("Jawaban objektif salah", "Menampilkan jawaban benar dan feedback singkat"),
        ("Esai ekuivalen secara numerik", "Diterima oleh pemeriksaan deterministik"),
        ("Esai memakai metode alternatif sah", "Dapat diterima melalui rubrik AI"),
        ("AI tidak tersedia", "Memberi pesan transparan; tidak mengklaim penilaian alternatif"),
        ("Hasil minimal 7/10", "Level lulus, progres tersimpan, level berikutnya terbuka"),
        ("Hasil kurang dari 7/10", "Heart berkurang satu dan kuis dapat diulang"),
    ], [3250, 6110])

    add_section(doc, "3.9 Perancangan Gamifikasi")
    add_body(doc, "Heart normal maksimum lima dan dapat bertambah hingga sepuluh melalui bonus misi. Jika heart di bawah lima, satu heart pulih setiap 30 menit. Streak dicatat ketika pengguna menyelesaikan satu sesi kuis, terlepas dari lulus atau belum. Streak bertambah jika penyelesaian berikutnya terjadi pada hari WIB yang berurutan, tetap jika masih pada hari yang sama, dan kembali menjadi satu jika terdapat jeda lebih dari satu hari. Rekor streak terlama dipertahankan.")
    add_body(doc, "Misi harian dipilih secara deterministik berdasarkan tanggal WIB dari beberapa target, misalnya menyelesaikan satu kuis, menjawab 10 atau 20 soal, atau lulus satu kuis. Ketika target tercapai untuk pertama kali pada hari tersebut, pengguna memperoleh lima heart. Desain ini menghubungkan hadiah dengan aktivitas belajar yang terukur.")

    add_section(doc, "3.10 Perancangan Basis Data")
    add_table(doc, ["Tabel", "Fungsi Utama"], [
        ("users", "Menyimpan akun, password hash, role, dan waktu pendaftaran"),
        ("ai_materi", "Menyimpan metadata materi dan lokasi file"),
        ("ai_soal", "Menyimpan pertanyaan, kunci, rubrik, tingkat, tipe, opsi, dan status aktif"),
        ("ai_auth_tokens", "Menyimpan hash token dan waktu kedaluwarsa"),
        ("ai_user_progress", "Menyimpan snapshot progres per pengguna"),
        ("ai_aktivitas", "Menyimpan log aktivitas yang diperlukan dashboard"),
        ("ai_chat_sessions/messages", "Menyimpan sesi dan pesan AskMatheal"),
        ("ai_password_reset_requests", "Menyimpan permintaan reset password"),
        ("ai_rate_limits", "Menyimpan pembatasan frekuensi permintaan"),
        ("ai_developers/meta", "Menyimpan informasi developer untuk landing page"),
    ], [2800, 6560])

    add_section(doc, "3.11 Implementasi Antarmuka User")
    add_body(doc, "Landing page menjelaskan manfaat, fitur, alur kerja, dan developer. Halaman login menyatukan ilustrasi \"Tahukah Kamu\" dengan card autentikasi. Homepage user menampilkan peta materi bergaya level, progres keseluruhan, streak, jumlah level lulus, heart, dan misi harian. Pada layar mobile, kartu progres ditempatkan di bagian atas agar informasi utama terlihat sebelum daftar materi.")
    add_body(doc, "Halaman materi menampilkan modul yang selaras dengan katalog kuis. Halaman kuis menggunakan tampilan gamifikasi dengan indikator nomor soal, progres, tipe soal, heart, feedback benar/salah, dan hasil akhir. Profil menampilkan poin, streak aktif dan terlama, jumlah heart, level lulus, progres mudah-sedang-sulit, grafik akumulasi soal berdasarkan hari WIB, serta catatan belajar.")

    add_section(doc, "3.12 Implementasi Antarmuka Admin")
    add_body(doc, "Role admin memiliki dashboard, pengguna, materi, bank soal, video panduan, aktivitas, dan profil. Dashboard mengambil snapshot ringkasan serta feed aktivitas. Data yang ditampilkan meliputi jumlah user, user aktif, aktivitas tujuh hari, materi dibuka, kuis diselesaikan, tren harian, status layanan, dan aktivitas terbaru. Penyegaran dapat dilakukan otomatis dan manual.")
    add_body(doc, "Bank soal dikelompokkan berdasarkan materi, tingkat, dan tipe. Validasi mencegah instruksi esai masuk ke soal objektif dan memastikan pilihan ganda memiliki opsi. Pengelolaan materi menggunakan file PDF dan membatasi path agar file di luar folder materi tidak dapat diakses melalui endpoint.")

    add_section(doc, "3.13 Implementasi AI Tutor dan Penilaian Esai")
    add_body(doc, "Server AI menggunakan Express, Axios, koneksi HTTPS persisten, timeout, percobaan ulang untuk gangguan sementara, pembatasan 60 permintaan per menit per alamat, serta pemeriksaan token. Prompt AskMatheal menetapkan peran tutor matematika berbahasa Indonesia, meminta langkah penyelesaian rinci, format Markdown, notasi LaTeX, dan kesimpulan akhir.")
    add_body(doc, "Prompt penilaian esai meminta keluaran JSON dengan score, answer_correct, method_valid, critical_error, dan feedback. Keputusan akhir tidak langsung mengikuti skor model, tetapi dihitung aplikasi: benar hanya jika jawaban akhir benar, metode valid, skor minimal 70, dan tidak terdapat kesalahan kritis. Pendekatan ini meningkatkan keterlacakan aturan keputusan dibanding sekadar meminta model menjawab \"benar\" atau \"salah\".")

    add_section(doc, "3.14 Keamanan Implementasi")
    add_bullets(doc, [
        "Password pengguna disimpan menggunakan fungsi password_hash dan diverifikasi menggunakan password_verify.",
        "Operasi basis data penting menggunakan prepared statement.",
        "Token disimpan dalam bentuk hash dan memiliki masa kedaluwarsa.",
        "Endpoint admin memeriksa role; endpoint AI memeriksa token melalui jalur internal.",
        "Origin produksi, kredensial database, kunci registrasi admin, internal API key, dan Gemini key diletakkan dalam environment server.",
        "Unggahan materi dibatasi pada PDF dan path file divalidasi.",
        "Aplikasi menyediakan rate limiting dan header keamanan dasar.",
    ])

    add_section(doc, "3.15 Rancangan Pengujian Black-Box")
    add_body(doc, "Black-box testing memeriksa kesesuaian keluaran terhadap input tanpa menilai struktur internal kode. Tabel berikut merupakan rancangan kasus uji. Kolom hasil aktual harus diisi selama pelaksanaan dan dilengkapi bukti.")
    test_rows = [
        ("Login user valid", "Masuk ke /home dan token tersimpan", "[Isi]"),
        ("Login salah", "Pesan kesalahan tampil tanpa membuat sesi", "[Isi]"),
        ("Akses role tidak sesuai", "Pengguna dialihkan ke halaman role yang benar", "[Isi]"),
        ("Membuka materi", "PDF sesuai topik tampil dan aktivitas tercatat", "[Isi]"),
        ("Memulai level terkunci", "Sistem menolak dan meminta menyelesaikan level sebelumnya", "[Isi]"),
        ("Kuis 10 soal", "Tiga mode soal muncul dan urutan diacak", "[Isi]"),
        ("Jawaban salah", "Jawaban benar dan feedback tampil lalu dapat dilanjutkan", "[Isi]"),
        ("Esai metode alternatif", "AI menilai proses, jawaban akhir, dan metode", "[Isi]"),
        ("Lulus minimal 7/10", "Progres tersimpan dan level berikutnya terbuka", "[Isi]"),
        ("Tidak lulus", "Heart berkurang dan kuis dapat diulang", "[Isi]"),
        ("Misi tercapai", "Bonus lima heart diberikan satu kali", "[Isi]"),
        ("Tiga level selesai", "Opsi sertifikat materi tersedia", "[Isi]"),
        ("Logout-login", "Progres akun dipulihkan", "[Isi]"),
        ("Dashboard admin", "Aktivitas user dan tren ditampilkan", "[Isi]"),
    ]
    add_table(doc, ["Skenario", "Hasil yang Diharapkan", "Hasil Aktual"], test_rows, [2700, 5000, 1660])
    add_placeholder(doc, "Jalankan setiap kasus uji, isi hasil aktual dengan Berhasil/Gagal, tambahkan tanggal, versi aplikasi, perangkat/browser, dan bukti tangkapan layar.")

    add_section(doc, "3.16 Pengujian Responsif dan Kompatibilitas")
    add_body(doc, "Pengujian responsif dilakukan pada beberapa lebar viewport yang mewakili mobile, tablet, laptop, dan desktop. Hal yang diperiksa meliputi tidak adanya overflow horizontal, keterbacaan font, urutan konten, ukuran target sentuh, grid kartu, modal, tabel admin, dan tampilan kuis. Browser yang disarankan adalah Chrome, Edge, dan Firefox versi yang digunakan responden.")
    add_table(doc, ["Viewport", "Representasi", "Pemeriksaan Utama", "Status"], [
        ("390 × 844", "Mobile", "Progress homepage di atas; kartu satu kolom; menu mobile", "[Isi]"),
        ("768 × 900", "Tablet", "Grid adaptif; modal dan grafik terbaca", "[Isi]"),
        ("1024 × 768", "Laptop kecil", "Sidebar dan konten tidak bertumpuk", "[Isi]"),
        ("1366 × 768", "Laptop umum", "Main content dan panel samping seimbang", "[Isi]"),
        ("1920 × 1080", "Desktop", "Lebar konten tetap nyaman dan tidak terlalu melebar", "[Isi]"),
    ], [1600, 1700, 4400, 1660])

    add_section(doc, "3.17 Pengujian Usability dengan SUS")
    add_body(doc, "Responden diminta menyelesaikan tugas: login, membuka materi, memulai kuis, menyelesaikan satu sesi, membaca progres profil, menggunakan AskMatheal, dan logout. Setelah itu responden mengisi 10 item SUS dalam skala 1 (sangat tidak setuju) sampai 5 (sangat setuju). Rumus skor: jumlah [(jawaban item ganjil - 1) + (5 - jawaban item genap)] × 2,5.")
    add_placeholder(doc, "Masukkan jumlah responden, skor setiap responden, nilai rata-rata SUS, median, rentang, serta interpretasi yang digunakan. Jangan menyimpulkan aplikasi usable sebelum data tersedia.")

    add_section(doc, "3.18 Validasi Materi, Soal, dan AI")
    add_body(doc, "Validasi akademik dilakukan oleh dosen atau ahli materi terhadap kesesuaian tujuan, kejelasan soal, kebenaran kunci, keseimbangan tingkat kesulitan, serta kualitas rubrik. Untuk AI, siapkan kumpulan jawaban uji: benar dengan metode kunci, benar dengan metode alternatif, salah tanda, salah konsep, jawaban akhir benar tetapi proses tidak valid, dan prompt injection. Hasil AI dibandingkan dengan keputusan ahli menggunakan persentase kesepakatan atau Cohen's kappa jika jumlah data memadai.")
    add_placeholder(doc, "Isi identitas dan kualifikasi validator, jumlah butir yang dinilai, skala penilaian, hasil validasi, revisi yang dilakukan, dan ukuran kesepakatan AI-dosen.")

    add_section(doc, "3.19 Teknik Analisis Data")
    add_numbers(doc, [
        "Persentase keberhasilan black-box = (jumlah kasus berhasil / jumlah seluruh kasus) × 100%.",
        "Skor SUS setiap responden dihitung pada rentang 0-100, kemudian diringkas dengan rata-rata dan sebaran.",
        "Validasi ahli dapat dihitung sebagai persentase kelayakan sesuai skala instrumen yang dipilih.",
        "Kesepakatan keputusan AI dan dosen dihitung dari confusion matrix: true positive, true negative, false positive, dan false negative.",
        "Jika ditambahkan pretest-posttest, gunakan uji normalitas dan uji beda yang sesuai serta laporkan effect size; jangan hanya membandingkan rata-rata secara deskriptif.",
    ])

    add_section(doc, "3.20 Ringkasan Implementasi Saat Ini")
    add_body(doc, "Berdasarkan pemeriksaan struktur proyek, komponen utama MatHeal telah tersedia: frontend React, API PHP/MySQL, server Node.js untuk Gemini, 17 berkas materi, katalog kuis, autentikasi berbasis token, progres pengguna, gamifikasi, sertifikat, dan dashboard admin. Namun, kesiapan ilmiah tetap memerlukan pengujian yang terdokumentasi. Dengan kata lain, keberadaan fitur merupakan hasil pengembangan, sedangkan efektivitas dan usability merupakan hasil penelitian yang harus dibuktikan melalui data.")


def chapter_four(doc):
    add_chapter(doc, "BAB IV\nPENUTUP")
    add_section(doc, "4.1 Kesimpulan")
    add_body(doc, "Berdasarkan proses analisis, perancangan, dan implementasi, telah dikembangkan aplikasi e-learning matematika informatika berbasis web bernama MatHeal. Aplikasi mengintegrasikan 17 materi, kuis tiga tingkat, tiga tipe soal, alur pembukaan level, umpan balik per soal, pencatatan progres, sertifikat per materi, AI Tutor, penilaian esai berbasis rubrik, serta dashboard administrator dalam satu sistem.")
    add_body(doc, "Integrasi Google Gemini pada MatHeal memiliki dua fungsi utama. Pertama, AskMatheal membantu pengguna memperoleh penjelasan matematika dalam bahasa Indonesia. Kedua, layanan penilaian esai memeriksa jawaban akhir dan validitas metode sehingga cara alternatif yang sah tidak harus sama persis dengan kunci. Keputusan akhir dibatasi oleh aturan eksplisit berupa jawaban benar, metode valid, skor minimal 70, dan tidak adanya kesalahan kritis. Mekanisme deterministik, autentikasi token, dan fallback ketika AI tidak tersedia melengkapi penggunaan model generatif.")
    add_body(doc, "Gamifikasi diterapkan melalui level mudah-sedang-sulit, poin, streak, rekor streak, heart, misi harian, progres, dan sertifikat. Elemen tersebut diarahkan untuk mendukung urutan dan konsistensi belajar. Kuis berisi 10 soal acak dan menyajikan jawaban yang benar ketika pengguna salah, sehingga sesi dapat diselesaikan dan hasil akhirnya menampilkan jumlah benar serta salah.")
    add_body(doc, "Dari sudut pengembangan perangkat lunak, MatHeal menggunakan arsitektur frontend React, backend PHP/MySQL, dan layanan AI Node.js/Express. Pemisahan tersebut mendukung pengelolaan fitur serta perlindungan kunci Gemini di server. Desain antarmuka juga telah dirancang responsif untuk perbedaan ukuran layar.")
    add_body(doc, "Kesimpulan tentang tingkat usability, kelayakan materi, akurasi penilaian AI, atau peningkatan hasil belajar belum boleh dinyatakan hanya berdasarkan implementasi. Kesimpulan kuantitatif final harus diperbarui setelah black-box testing, uji responsif, SUS, validasi ahli, dan—jika ingin mengukur efektivitas belajar—pretest-posttest selesai dilakukan.")
    add_placeholder(doc, "Setelah pengujian, tambahkan satu paragraf yang memuat angka faktual: persentase black-box, rata-rata SUS, hasil validasi ahli, akurasi/kesepakatan AI, dan hasil statistik belajar bila diuji.")

    add_section(doc, "4.2 Keterbatasan")
    add_bullets(doc, [
        "Kualitas penilaian esai bergantung pada ketersediaan layanan Gemini, kualitas prompt, kunci, dan rubrik.",
        "AI generatif masih dapat menghasilkan penilaian yang tidak konsisten; keputusan berisiko tinggi tetap memerlukan verifikasi dosen.",
        "Sebagian state interaksi disimpan pada browser dan disinkronkan ke akun, sehingga perlu pengujian lebih lanjut pada banyak perangkat dan konflik sinkronisasi.",
        "Bank soal dan materi memerlukan validasi serta pembaruan berkala agar tetap jelas, tidak rancu, dan seimbang.",
        "Dashboard saat ini berfungsi untuk ringkasan operasional, belum menjadi learning analytics adaptif yang mendalam.",
        "Tanpa desain eksperimen, penelitian ini belum dapat membuktikan hubungan kausal antara penggunaan MatHeal dan peningkatan hasil belajar.",
    ])

    add_section(doc, "4.3 Saran")
    add_numbers(doc, [
        "Melaksanakan uji pengguna dengan sampel yang terdokumentasi, memakai SUS dan observasi keberhasilan tugas.",
        "Menambahkan eksperimen pretest-posttest dengan kelompok pembanding apabila tujuan penelitian lanjutan adalah mengukur efektivitas pembelajaran.",
        "Membangun golden dataset jawaban esai yang telah dinilai dosen untuk mengukur akurasi, konsistensi, false positive, dan false negative AI.",
        "Menambahkan panel review dosen untuk mengoreksi keputusan AI serta menyempurnakan rubrik soal.",
        "Memindahkan seluruh state progres kritis ke basis data sebagai sumber utama, dengan localStorage hanya sebagai cache antarmuka.",
        "Menambahkan pengujian otomatis frontend, endpoint PHP, server Node, aturan heart-streak, pengacakan soal, dan alur autentikasi pada continuous integration.",
        "Meningkatkan aksesibilitas melalui navigasi keyboard, kontras, label pembaca layar, alternatif teks, dan pengujian WCAG.",
        "Mengembangkan learning analytics yang tetap menjaga privasi, misalnya tingkat penyelesaian materi dan pola kesalahan agregat tanpa menyimpan jawaban sensitif yang tidak diperlukan.",
        "Menyediakan kebijakan privasi, persetujuan penggunaan AI, mekanisme pelaporan jawaban AI yang salah, backup, monitoring, dan audit log sebelum penggunaan skala besar.",
        "Melakukan optimasi deployment, pengujian beban, observability, serta evaluasi biaya Gemini sebelum aplikasi digunakan oleh banyak pengguna.",
    ])


def bibliography(doc):
    add_chapter(doc, "DAFTAR PUSTAKA")
    refs = [
        "Bai, S., Hew, K. F., & Huang, B. (2020). Does gamification improve student learning outcome? Evidence from a meta-analysis and synthesis of qualitative data in educational contexts. Educational Research Review, 30, 100322. https://doi.org/10.1016/j.edurev.2020.100322",
        "Brooke, J. (1996). SUS: A quick and dirty usability scale. Dalam P. W. Jordan, B. Thomas, B. A. Weerdmeester, & A. L. McClelland (Ed.), Usability Evaluation in Industry (hlm. 189-194). Taylor & Francis.",
        "Deterding, S., Dixon, D., Khaled, R., & Nacke, L. (2011). From game design elements to gamefulness: Defining gamification. Proceedings of the 15th International Academic MindTrek Conference, 9-15. https://doi.org/10.1145/2181037.2181040",
        "Google. (2026). Gemini API: Generating content. Google AI for Developers. https://ai.google.dev/api/generate-content",
        "ISO/IEC. (2023). ISO/IEC 25010:2023 Systems and software engineering—Systems and software Quality Requirements and Evaluation (SQuaRE)—Product quality model. International Organization for Standardization. https://www.iso.org/standard/78176.html",
        "Miao, F., & Holmes, W. (2023). Guidance for generative AI in education and research. UNESCO. https://unesdoc.unesco.org/ark:/48223/pf0000386693",
        "Niño-Rojas, F., Lancheros-Cuesta, D., Jiménez-Valderrama, M. T. P., Mestre, G., & Gómez, S. (2024). Systematic review: Trends in intelligent tutoring systems in mathematics teaching and learning. International Journal of Education in Mathematics, Science and Technology, 12(1), 203-229.",
        "PHP Documentation Group. (2026). Prepared statements. PHP Manual. https://www.php.net/manual/en/mysqli.quickstart.prepared-statements.php",
        "React Team. (2026). Quick start: React. https://react.dev/learn",
        "Sailer, M., & Homner, L. (2020). The gamification of learning: A meta-analysis. Educational Psychology Review, 32, 77-112. https://doi.org/10.1007/s10648-019-09498-w",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-1.0)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(ref)
        set_run_font(r, size=11)


def set_core_properties(doc):
    props = doc.core_properties
    props.title = "Draf Penulisan Ilmiah MatHeal BAB I-IV"
    props.subject = "Pengembangan aplikasi e-learning matematika informatika berbasis web dengan AI Tutor dan gamifikasi"
    props.author = "Peneliti MatHeal"
    props.keywords = "MatHeal, e-learning, matematika informatika, AI Tutor, Google Gemini, gamifikasi"
    props.comments = "Draf akademik; lengkapi seluruh placeholder dengan data penelitian faktual."


def build():
    create_diagrams()
    doc = Document()
    configure_document(doc)
    set_core_properties(doc)
    add_cover(doc)
    add_front_matter(doc)
    chapter_one(doc)
    chapter_two(doc)
    chapter_three(doc)
    chapter_four(doc)
    bibliography(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
