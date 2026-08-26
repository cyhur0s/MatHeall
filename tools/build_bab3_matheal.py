"""Membangun dokumen BAB III MatHeal dari fakta implementasi proyek.

Dokumen sengaja mempertahankan placeholder pada hasil uji yang belum dilakukan.
"""
from pathlib import Path
import sys

from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_pi_draft as source


OUT = Path(__file__).resolve().parents[1] / "docs" / "Bab_III_Metodologi_dan_Pengembangan_MatHeal.docx"


def after(paragraph, text, style="Normal"):
    """Sisipkan paragraf Word sesudah paragraf tertentu."""
    new_p = source.OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.style = style
    run = inserted.add_run(text)
    source.set_run_font(run, size=12)
    return inserted


def find(doc, exact):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact:
            return paragraph
    raise ValueError(f"Paragraf tidak ditemukan: {exact}")


def replace_heading(doc, old, new):
    paragraph = find(doc, old)
    paragraph.text = new
    for run in paragraph.runs:
        source.set_run_font(run, size=13 if paragraph.style.name == "Heading 2" else 12, bold=True)
    return paragraph


def add_references(doc):
    doc.add_page_break()
    title = doc.add_paragraph()
    title.style = "Heading 1"
    run = title.add_run("DAFTAR PUSTAKA RUJUKAN BAB III")
    source.set_run_font(run, size=14, bold=True)
    refs = [
        "Balalle, H. (2024). Exploring student engagement in technology-based education in relation to gamification, online/distance learning, and other factors: A systematic literature review. Social Sciences & Humanities Open, 9, 100870. https://doi.org/10.1016/j.ssaho.2024.100870",
        "Brooke, J. (1996). SUS: A quick and dirty usability scale. In P. W. Jordan et al. (Eds.), Usability Evaluation in Industry (pp. 189-194). Taylor & Francis.",
        "ISO/IEC. (2023). ISO/IEC 25010:2023 Systems and software engineering - Systems and software Quality Requirements and Evaluation (SQuaRE) - Product quality model. International Organization for Standardization.",
        "Pressman, R. S., & Maxim, B. R. (2020). Software Engineering: A Practitioner's Approach (9th ed.). McGraw-Hill.",
        "Ruiz, J. J. R., Sanchez, A. D. V., & Figueredo, O. R. B. (2024). Impact of gamification on school engagement: A systematic review. Frontiers in Education, 9, 1466926. https://doi.org/10.3389/feduc.2024.1466926",
        "Sailer, M., & Homner, L. (2020). The gamification of learning: A meta-analysis. Educational Psychology Review, 32, 77-112. https://doi.org/10.1007/s10648-019-09498-w",
        "[Lengkapi dengan artikel AI Tutor/penilaian esai yang benar-benar digunakan dari Bab II dan sesuai gaya sitasi kampus.]",
    ]
    for item in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = source.Cm(1)
        p.paragraph_format.first_line_indent = source.Cm(-1)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        source.set_run_font(p.add_run(item), size=12)


def build():
    source.create_diagrams()
    doc = Document()
    source.configure_document(doc)
    source.set_core_properties(doc)
    source.chapter_three(doc)

    # Perbaikan struktur: masalah -> kebutuhan pengguna -> rancangan/implementasi.
    p_problem = replace_heading(doc, "3.5 Analisis Kebutuhan Pengguna", "3.5 Analisis Permasalahan")
    problem = after(
        p_problem,
        "Permasalahan yang menjadi dasar pengembangan MatHeal adalah materi matematika informatika dan latihan yang belum terhubung dalam satu alur belajar, umpan balik terhadap kesalahan yang terbatas, serta kesulitan pengguna dalam memantau kemajuan belajar. Di sisi lain, pembelajaran memerlukan latihan berulang yang terstruktur tanpa menjadikan penghargaan sebagai tujuan utama. Kebutuhan tersebut dijawab melalui materi, kuis bertingkat, umpan balik, progres, dan gamifikasi yang melekat pada aktivitas belajar. Temuan kajian mutakhir menunjukkan bahwa gamifikasi berpotensi mendukung keterlibatan, tetapi hasilnya bergantung pada rancangan, konteks, dan evaluasi yang memadai (Balalle, 2024; Ruiz et al., 2024).",
    )
    p_need_user = replace_heading(doc, "3.5.1 Kebutuhan Fungsional User", "3.6.1 Kebutuhan Fungsional User")
    p_need = after(problem, "3.6 Analisis Kebutuhan Pengguna", "Heading 2")
    # Pindahkan heading kebutuhan sebelum subbagian pengguna secara visual dengan XML.
    p_need._p.addnext(p_need_user._p)
    replace_heading(doc, "3.5.2 Kebutuhan Fungsional Admin", "3.6.2 Kebutuhan Fungsional Admin")
    replace_heading(doc, "3.5.3 Kebutuhan Nonfungsional", "3.6.3 Kebutuhan Nonfungsional")

    mapping = {
        "3.6 Arsitektur Sistem": "3.7 Arsitektur Sistem",
        "3.7 Perancangan Alur Belajar": "3.8 Perancangan Alur Belajar",
        "3.8 Perancangan Kuis": "3.9 Perancangan Kuis",
        "3.9 Perancangan Gamifikasi": "3.10 Perancangan Gamifikasi",
        "3.10 Perancangan Basis Data": "3.11 Perancangan Basis Data",
        "3.11 Implementasi Antarmuka User": "3.12 Implementasi Antarmuka User",
        "3.12 Implementasi Antarmuka Admin": "3.13 Implementasi Antarmuka Admin",
        "3.13 Implementasi AI Tutor dan Penilaian Esai": "3.14 Implementasi AI Tutor dan Penilaian Esai",
        "3.14 Keamanan Implementasi": "3.15 Keamanan Implementasi",
        "3.15 Rancangan Pengujian Black-Box": "3.16 Rancangan Pengujian Black-Box Testing",
        "3.16 Pengujian Responsif dan Kompatibilitas": "3.17 Pengujian Responsif, Kompatibilitas, dan Usability",
    }
    for old, new in mapping.items():
        replace_heading(doc, old, new)

    # SUS merupakan instrumen usability pada evaluasi produk; bukan hasil yang telah diklaim.
    sus_heading = find(doc, "3.17 Pengujian Usability dengan SUS")
    sus_heading._element.getparent().remove(sus_heading._element)
    sus_text = find(doc, "Responden diminta menyelesaikan tugas: login, membuka materi, memulai kuis, menyelesaikan satu sesi, membaca progres profil, menggunakan AskMatheal, dan logout. Setelah itu responden mengisi 10 item SUS dalam skala 1 (sangat tidak setuju) sampai 5 (sangat setuju). Rumus skor: jumlah [(jawaban item ganjil - 1) + (5 - jawaban item genap)] × 2,5.")
    sus_text.text = "Untuk evaluasi usability, responden menyelesaikan tugas login, membuka materi, memulai dan menyelesaikan kuis, membaca progres, menggunakan AskMatheal, lalu logout. Setelah itu responden mengisi System Usability Scale (SUS) dengan skala 1-5. Skor dihitung dengan rumus jumlah [(item ganjil - 1) + (5 - item genap)] × 2,5. Hasil SUS hanya diisi setelah pengujian pengguna dilakukan."
    for run in sus_text.runs:
        source.set_run_font(run, size=12)

    # Jaga penomoran akhir yang dikehendaki pengguna.
    replace_heading(doc, "3.18 Validasi Materi, Soal, dan AI", "3.18 Validasi Materi, Soal, dan AI")
    replace_heading(doc, "3.19 Teknik Analisis Data", "3.19 Teknik Analisis Data")
    replace_heading(doc, "3.20 Ringkasan Implementasi Saat Ini", "3.20 Ringkasan Implementasi Saat Ini")

    add_references(doc)
    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
