from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parents[1] / "docs" / "Revisi_Bab_I_Pendahuluan_MatHeal.docx"


def font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    font(run, 10)
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Cm(4)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    for name, size in (("Heading 1", 14), ("Heading 2", 12)):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
    add_page_number(section.footer.paragraphs[0])


def paragraph(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    font(p.add_run(text), 12)
    return p


def heading(doc, number, title):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    font(p.add_run(f"{number} {title}"), 12, bold=True)
    return p


def item(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(label + " ")
    font(r, 12, bold=True)
    font(p.add_run(text), 12)
    return p


def build():
    doc = Document()
    setup(doc)
    props = doc.core_properties
    props.title = "Revisi Bab I Pendahuluan MatHeal"
    props.subject = "Penulisan ilmiah pengembangan MatHeal"
    props.author = "Peneliti MatHeal"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    font(title.add_run("BAB I\nPENDAHULUAN"), 14, bold=True)

    heading(doc, "1.1", "Latar Belakang")
    paragraph(doc, "Matematika informatika merupakan salah satu fondasi penting dalam pendidikan komputasi. Konsep seperti logika, himpunan, relasi, fungsi, bilangan biner, matriks, peluang, dan graf diperlukan untuk memahami algoritma, struktur data, basis data, serta mata kuliah komputasi lanjutan. Oleh karena itu, mahasiswa perlu memperoleh materi yang terstruktur, contoh penyelesaian bertahap, latihan yang bervariasi, dan umpan balik atas kesalahan yang dilakukan selama belajar.")
    paragraph(doc, "Kebutuhan tersebut belum selalu terpenuhi melalui media belajar yang hanya menyajikan dokumen atau kuis dengan pencocokan jawaban akhir. Pada materi matematika, jawaban yang benar dapat diperoleh melalui notasi dan metode penyelesaian yang berbeda. Sistem yang hanya memeriksa kesamaan teks berisiko menolak jawaban ekuivalen, sedangkan pengguna yang mengalami kesulitan memerlukan penjelasan konsep dan langkah penyelesaian yang lebih terarah. Kondisi ini menunjukkan perlunya media pembelajaran yang menghubungkan materi, latihan, umpan balik, dan pemantauan kemajuan dalam satu alur.")
    paragraph(doc, "E-learning berbasis web dapat menyediakan fleksibilitas akses terhadap materi dan latihan. Selain itu, elemen gamifikasi seperti level, poin, progres, streak, heart, misi harian, dan sertifikat dapat digunakan untuk memberi tujuan belajar jangka pendek serta umpan balik atas aktivitas pengguna. Penerapannya perlu dikaitkan dengan aktivitas belajar, bukan hanya sebagai unsur visual, agar tetap mendukung keteraturan latihan dan pencapaian kompetensi (Deterding dkk., 2011; Ruiz dkk., 2024).")
    paragraph(doc, "Pemanfaatan AI generatif juga membuka peluang untuk menyediakan tanya jawab dan umpan balik otomatis. Namun, AI Tutor tidak dapat diperlakukan sebagai satu-satunya sumber kebenaran akademik. Penilaian esai perlu mempertimbangkan ketepatan jawaban akhir, validitas metode, serta adanya kesalahan kritis. Oleh sebab itu, AI pada MatHeal dirancang sebagai pendamping belajar yang memberikan penjelasan dan umpan balik berdasarkan prompt serta rubrik terstruktur, bukan sebagai pengganti pengajar.")
    paragraph(doc, "Berdasarkan telaah fitur aplikasi e-learning dan kebutuhan alur belajar matematika informatika, peneliti belum menemukan penerapan lokal yang secara terintegrasi menghubungkan modul materi, kuis bertingkat, progres belajar, gamifikasi multi-elemen, AI Tutor, serta penilaian esai yang mempertimbangkan metode alternatif. Kesenjangan tersebut menjadi dasar pengembangan MatHeal, yaitu aplikasi e-learning matematika informatika berbasis web yang dirancang untuk mendukung proses belajar mandiri secara terstruktur.")
    paragraph(doc, "MatHeal menyediakan materi dalam bentuk modul, kuis dengan tingkat mudah, sedang, dan sulit, serta umpan balik pada jawaban pengguna. Sistem menerapkan poin, streak, heart, misi harian, indikator progres, dan sertifikat sebagai dukungan konsistensi belajar. Fitur AskMatheal menggunakan Google Gemini untuk menjawab pertanyaan matematika dan membantu menilai jawaban esai berdasarkan rubrik. Penelitian ini berfokus pada pengembangan dan pemeriksaan fungsi aplikasi; penelitian tidak mengukur peningkatan hasil belajar, usability pengguna, maupun kesepakatan penilaian AI dengan dosen karena tidak melibatkan responden atau validator ahli.")

    heading(doc, "1.2", "Rumusan Masalah")
    paragraph(doc, "Berdasarkan latar belakang tersebut, rumusan masalah dalam penelitian ini adalah sebagai berikut.", indent=False)
    item(doc, "1.", "Bagaimana merancang dan mengembangkan aplikasi e-learning matematika informatika berbasis web yang mengintegrasikan materi, kuis bertingkat, progres belajar, sertifikat, dan dashboard administrator?")
    item(doc, "2.", "Bagaimana mengintegrasikan AI Tutor Google Gemini sebagai media tanya jawab matematika dan penilai jawaban esai berbasis rubrik yang mempertimbangkan jawaban akhir serta validitas metode penyelesaian?")
    item(doc, "3.", "Bagaimana menerapkan gamifikasi berupa level, poin, streak, heart, misi harian, progres, dan sertifikat untuk mendukung keteraturan belajar pengguna?")
    item(doc, "4.", "Bagaimana menguji kesesuaian fungsi utama dan tampilan responsif aplikasi MatHeal melalui skenario black-box testing yang dilakukan oleh peneliti?")

    heading(doc, "1.3", "Batasan Masalah")
    paragraph(doc, "Agar penelitian tetap terarah, ruang lingkup pengembangan MatHeal dibatasi sebagai berikut.", indent=False)
    item(doc, "1.", "Materi dibatasi pada 17 topik matematika informatika yang tersedia pada aplikasi saat penelitian dilakukan.")
    item(doc, "2.", "Aplikasi memiliki dua peran utama, yaitu user sebagai pengguna pembelajaran dan admin sebagai pengelola materi, soal, pengguna, serta aktivitas aplikasi.")
    item(doc, "3.", "Kuis memiliki tiga tingkat kesulitan, yaitu mudah, sedang, dan sulit. Level berikutnya dibuka setelah pengguna lulus level sebelumnya dengan nilai minimal 70 persen.")
    item(doc, "4.", "Setiap sesi kuis terdiri atas 10 soal dari bank soal yang mencakup pilihan ganda, benar-salah, dan esai.")
    item(doc, "5.", "AI Tutor menggunakan Google Gemini melalui layanan backend. Respons AI diposisikan sebagai umpan balik otomatis dan tidak menggantikan keputusan akademik pengajar.")
    item(doc, "6.", "Sertifikat merupakan bukti penyelesaian materi pada aplikasi dan bukan sertifikat akademik resmi.")
    item(doc, "7.", "Penelitian tidak melibatkan responden mahasiswa, kuesioner System Usability Scale, validasi materi oleh dosen, maupun pengukuran akurasi AI terhadap penilaian ahli.")
    item(doc, "8.", "Evaluasi fungsional dilakukan melalui skenario black-box testing oleh peneliti setelah implementasi aplikasi siap diuji. Efektivitas pembelajaran tidak diukur melalui pretest-posttest pada penelitian ini.")

    heading(doc, "1.4", "Tujuan Penelitian")
    paragraph(doc, "Tujuan penelitian ini adalah sebagai berikut.", indent=False)
    item(doc, "1.", "Mengembangkan aplikasi e-learning matematika informatika berbasis web bernama MatHeal yang menghubungkan materi, kuis bertingkat, progres, sertifikat, dan dashboard administrator.")
    item(doc, "2.", "Mengintegrasikan AI Tutor Google Gemini untuk memberi penjelasan matematika dan menilai jawaban esai berdasarkan ketepatan jawaban akhir serta validitas proses penyelesaian.")
    item(doc, "3.", "Menerapkan gamifikasi berupa level, poin, streak, heart, misi harian, progres, dan sertifikat untuk mendukung keteraturan belajar pengguna.")
    item(doc, "4.", "Menguji kesesuaian fungsi utama dan tampilan responsif MatHeal melalui skenario black-box testing yang dilaksanakan oleh peneliti.")

    heading(doc, "1.5", "Manfaat Penelitian")
    paragraph(doc, "Manfaat penelitian ini terdiri atas manfaat praktis dan manfaat akademik.", indent=False)
    item(doc, "1.", "Bagi pengguna, MatHeal diharapkan menjadi media belajar pendamping yang menyediakan materi, latihan bertingkat, umpan balik, dan pemantauan progres dalam satu aplikasi.")
    item(doc, "2.", "Bagi administrator, aplikasi menyediakan fasilitas untuk mengelola materi, bank soal, pengguna, serta aktivitas pembelajaran.")
    item(doc, "3.", "Bagi peneliti berikutnya, dokumentasi pengembangan dan rancangan pengujian dapat digunakan sebagai dasar pengembangan evaluasi usability, validasi akademik AI, atau pengukuran efektivitas pembelajaran.")

    heading(doc, "1.6", "Metode Penelitian")
    paragraph(doc, "Penelitian ini menggunakan pendekatan Research and Development (R&D) dengan model pengembangan perangkat lunak prototyping. Model ini dipilih karena rancangan fitur, alur pembelajaran, dan antarmuka MatHeal dikembangkan secara bertahap melalui proses identifikasi kebutuhan, perancangan cepat, pembangunan prototipe, evaluasi internal, penyempurnaan, dan pengujian fungsional. Pendekatan ini sesuai untuk menghasilkan produk perangkat lunak pendidikan yang dapat diperbaiki secara iteratif berdasarkan temuan selama pengembangan (Pressman & Maxim, 2020).")
    paragraph(doc, "Data pengembangan diperoleh melalui studi pustaka, analisis kebutuhan aplikasi, dan dokumentasi implementasi. Pengujian fungsional dirancang menggunakan black-box testing yang dilakukan oleh peneliti untuk membandingkan keluaran aplikasi dengan hasil yang diharapkan pada skenario utama, seperti autentikasi, akses materi, kuis, progres, gamifikasi, AI Tutor, sertifikat, dan dashboard admin. Penelitian ini tidak menggunakan SUS maupun pengujian yang memerlukan responden atau dosen validator. Penjelasan rinci mengenai tahapan prototipe, kebutuhan sistem, implementasi, dan rancangan pengujian disajikan pada Bab III.")

    heading(doc, "1.7", "Sistematika Penulisan")
    paragraph(doc, "Sistematika penulisan disusun dalam empat bab sebagai berikut.", indent=False)
    item(doc, "BAB I", "Pendahuluan, berisi latar belakang, rumusan masalah, batasan masalah, tujuan penelitian, manfaat penelitian, metode penelitian, dan sistematika penulisan.")
    item(doc, "BAB II", "Landasan Teori, berisi teori dan penelitian terdahulu yang berkaitan dengan e-learning, matematika informatika, gamifikasi, AI Tutor, penilaian esai, teknologi web, keamanan dasar, dan pengujian perangkat lunak.")
    item(doc, "BAB III", "Analisis, Perancangan, Implementasi, dan Pengujian, berisi analisis permasalahan dan kebutuhan, tahapan pengembangan prototipe, arsitektur sistem, perancangan basis data, implementasi antarmuka, integrasi AI Tutor, keamanan, serta rancangan dan hasil pengujian jika telah dilaksanakan.")
    item(doc, "BAB IV", "Penutup, berisi kesimpulan berdasarkan hasil pengembangan dan pengujian yang telah dilakukan, keterbatasan penelitian, serta saran pengembangan berikutnya.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
