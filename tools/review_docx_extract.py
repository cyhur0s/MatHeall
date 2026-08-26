from docx import Document
import sys

sys.stdout.reconfigure(encoding="utf-8")

p = r"D:\KULIAH\Penelitian Ilmiah\testing penulisan.docx"
d = Document(p)
print("PARAGRAPHS", len(d.paragraphs), "TABLES", len(d.tables))
for i, x in enumerate(d.paragraphs, 1):
    if not 80 <= i <= 165:
        continue
    text = x.text.strip()
    if text:
        print(f"P{i:03d}|{x.style.name}|{text}")
for ti, table in enumerate(d.tables, 1):
    print(f"\nTABLE {ti}: {len(table.rows)}x{len(table.columns)}")
    for row in table.rows:
        print(" | ".join(cell.text.replace("\n", " / ") for cell in row.cells))
